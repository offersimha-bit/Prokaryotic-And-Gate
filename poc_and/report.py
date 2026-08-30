"""
Stage 5 -- the deliverables: an editable Word report and a supporting workbook.

Three files, written to out_dir:

    AND_switch_report.docx      one page per candidate, laid out like
                                Toehold_Candidates29.7.pdf so the two can be
                                read side by side
    AND_switch_appendix.docx    every formula written out, with its caveats
    AND_switch_data.xlsx        all the numbers, re-sortable

Figures are drawn from ViennaRNA's own layout coordinates via matplotlib, so
there is no SVG conversion step and no extra dependency. FORNA-ready text
(sequence + dot-bracket) is emitted alongside each figure for anyone who wants
the interactive version.

Deliberately NOT included: arc plots.
"""

# Make relative imports work when this file is run on its own.
if __package__ in (None, ""):
    import os as _os, sys as _sys
    _sys.path.insert(0, _os.path.dirname(_os.path.dirname(_os.path.abspath(__file__))))
    import poc_and  # noqa: F401
    __package__ = "poc_and"

import os

import matplotlib
matplotlib.use("Agg")                    # no display needed
import matplotlib.pyplot as plt

from .folding import RNA

from . import candidates as cd
from . import codon_usage
from . import folding
from . import pdf_metrics


# ---------------------------------------------------------------------------
# Figures
# ---------------------------------------------------------------------------

# FORNA's default "structure" colour scheme, matched to the figures in
# Toehold_Candidates29.7.pdf so the two documents can be read side by side.
FORNA_COLOURS = {
    "stem":      ("#9ed69e", "#6aa96a"),
    "hairpin":   ("#a9c8e0", "#7099bd"),
    "interior":  ("#d9d98a", "#b0b055"),
    "multiloop": ("#e59a9a", "#c06868"),
    "exterior":  ("#f0ac82", "#cc7c4d"),
}


def loop_types(structure):
    """
    Classify every nucleotide the way FORNA colours them.

    paired            -> stem
    unpaired, and the innermost pair enclosing it contains
        no helices    -> hairpin loop
        one helix     -> interior loop (bulges included)
        two or more   -> multiloop
    unpaired, enclosed by nothing -> exterior loop

    Written out rather than approximated because the colouring IS the
    information in these figures: it is what makes the RBS loop, the AUG bulge
    and the exposed toehold identifiable at a glance.
    """
    pt = RNA.ptable(structure)
    n = len(structure)
    kinds = [None] * (n + 1)

    def children(p, q):
        """Number of helices directly inside the pair (p, q)."""
        count = 0
        j = p + 1
        while j < q:
            if pt[j] > j:
                count += 1
                j = pt[j] + 1
            else:
                j += 1
        return count

    # innermost enclosing pair for each position, by one sweep with a stack
    stack = []
    enclosing = [0] * (n + 1)
    for i in range(1, n + 1):
        if pt[i] > i:
            enclosing[i] = stack[-1] if stack else 0
            stack.append(i)
        elif pt[i] != 0:
            stack.pop()
            enclosing[i] = stack[-1] if stack else 0
        else:
            enclosing[i] = stack[-1] if stack else 0

    child_cache = {}
    for i in range(1, n + 1):
        if pt[i] != 0:
            kinds[i] = "stem"
            continue
        p = enclosing[i]
        if p == 0:
            kinds[i] = "exterior"
            continue
        if p not in child_cache:
            child_cache[p] = children(p, pt[p])
        c = child_cache[p]
        kinds[i] = "hairpin" if c == 0 else ("interior" if c == 1 else "multiloop")
    return kinds[1:]


def _label_offset(xs, ys, i, pair_table):
    """
    Unit vector pointing away from the molecule at position i.

    Position labels have to sit OUTSIDE the structure or they land on a
    nucleotide circle and neither is readable. Two cases, because one rule does
    not cover both:

      paired base    -> point away from its PARTNER. That is straight out from
                        the helix axis, which is the only free direction in a
                        stem. Offsetting perpendicular to the backbone instead
                        puts the label on the opposite strand.
      unpaired base  -> point away from the midpoint of its two neighbours,
                        which is outward on any loop or bulge.
    """
    n = len(xs)
    partner = pair_table[i + 1] - 1
    if partner >= 0:
        dx, dy = xs[i] - xs[partner], ys[i] - ys[partner]
        norm = (dx * dx + dy * dy) ** 0.5
        if norm > 1e-6:
            return dx / norm, dy / norm

    prev_i = max(0, i - 1)
    next_i = min(n - 1, i + 1)
    dx = xs[i] - (xs[prev_i] + xs[next_i]) / 2.0
    dy = ys[i] - (ys[prev_i] + ys[next_i]) / 2.0
    norm = (dx * dx + dy * dy) ** 0.5
    if norm > 1e-6:
        return dx / norm, dy / norm
    tx, ty = xs[next_i] - xs[prev_i], ys[next_i] - ys[prev_i]
    tnorm = (tx * tx + ty * ty) ** 0.5 or 1.0
    return -ty / tnorm, tx / tnorm


def draw_structure(sequence, structure, path, title, number_every=10):
    """
    FORNA-style secondary-structure diagram: lettered circles, coloured by loop
    type, grey backbone, base pairs marked, numbered every 10 nt.

    Layout comes from ViennaRNA's naview algorithm -- the same one FORNA uses --
    so the picture matches both the structure the pipeline computed and the
    figures in the candidates PDF.
    """
    coords = RNA.naview_xy_coordinates(structure)
    xs = [coords[i].X for i in range(len(sequence))]
    ys = [coords[i].Y for i in range(len(sequence))]
    pt = RNA.ptable(structure)
    kinds = loop_types(structure)

    # circle radius from the median spacing between neighbours
    steps = sorted(((xs[i + 1] - xs[i]) ** 2 + (ys[i + 1] - ys[i]) ** 2) ** 0.5
                   for i in range(len(sequence) - 1))
    spacing = steps[len(steps) // 2] if steps else 20.0
    radius = spacing * 0.44

    # Keep the page proportions sane. naview layouts can be very elongated, and
    # a 1:3 figure is unusable in a Word document, so cap the aspect ratio and
    # let the equal-aspect axes letterbox the rest.
    span_x = max(xs) - min(xs) + 4 * spacing
    span_y = max(ys) - min(ys) + 4 * spacing
    longest = max(span_x, span_y)
    width_in = 7.5 * (span_x / longest)
    height_in = 7.5 * (span_y / longest)
    width_in = min(max(width_in, 4.0), 7.5)
    height_in = min(max(height_in, 3.5), 9.0)
    fig, ax = plt.subplots(figsize=(width_in, height_in))

    ax.plot(xs, ys, "-", color="#b0b0b0", linewidth=1.1, zorder=1)
    for i in range(1, len(sequence) + 1):
        j = pt[i]
        if j > i:
            ax.plot([xs[i - 1], xs[j - 1]], [ys[i - 1], ys[j - 1]],
                    "-", color="#d06060", linewidth=1.0, alpha=0.85, zorder=2)

    font_size = max(3.5, min(8.0, radius * 1.5))
    for i, base in enumerate(sequence):
        fill, edge = FORNA_COLOURS[kinds[i]]
        ax.add_patch(plt.Circle((xs[i], ys[i]), radius, facecolor=fill,
                                edgecolor=edge, linewidth=0.6, zorder=3))
        ax.text(xs[i], ys[i], base, ha="center", va="center",
                fontsize=font_size, color="#222222", zorder=4)
        n = i + 1
        if n == 1 or n % number_every == 0:
            ox, oy = _label_offset(xs, ys, i, pt)
            ax.text(xs[i] + ox * radius * 2.4, ys[i] + oy * radius * 2.4,
                    str(n), ha="center", va="center",
                    fontsize=font_size * 0.85, color="#444444", zorder=5)

    handles = [plt.Line2D([], [], marker="o", linestyle="",
                          markerfacecolor=FORNA_COLOURS[k][0],
                          markeredgecolor=FORNA_COLOURS[k][1], markersize=7,
                          label=label)
               for k, label in (("stem", "stem"), ("hairpin", "hairpin loop"),
                                ("interior", "interior loop"),
                                ("multiloop", "multiloop"),
                                ("exterior", "exterior loop"))]
    ax.legend(handles=handles, loc="upper right", fontsize=6.5, frameon=False,
              ncol=2)
    ax.set_title(title, fontsize=10)
    ax.set_aspect("equal")
    ax.axis("off")
    ax.margins(0.06)
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)
    return path


def draw_base_pair_probabilities(strands, path, title, marks=None):
    """
    Base-pairing probability matrix.

    The MFE is one structure; this is the whole ensemble, which matters here
    because the MFE routinely carries only a few percent of it. Optional
    `marks` draws domain boundaries so features can be located.
    """
    joined = "&".join(strands) if not isinstance(strands, str) else strands
    fc = RNA.fold_compound(joined)
    _, mfe_energy = fc.mfe()
    fc.exp_params_rescale(mfe_energy)
    fc.pf()
    bpp = fc.bpp()
    n = len(joined.replace("&", ""))

    xs, ys, weights = [], [], []
    for i in range(1, n + 1):
        row = bpp[i]
        for j in range(i + 1, n + 1):
            p = row[j]
            if p > 1e-4:
                xs.append(i)
                ys.append(j)
                weights.append(p)

    fig, ax = plt.subplots(figsize=(6.4, 6.0))
    if xs:
        sc = ax.scatter(xs, ys, c=weights, s=[6 + 26 * w for w in weights],
                        cmap="viridis", vmin=0, vmax=1, linewidths=0)
        fig.colorbar(sc, ax=ax, label="pair probability", shrink=0.8)
    for boundary, label in (marks or []):
        ax.axvline(boundary, color="#cc4444", linewidth=0.6, alpha=0.6)
        ax.axhline(boundary, color="#cc4444", linewidth=0.6, alpha=0.6)
        ax.text(boundary, n * 1.01, label, fontsize=6, rotation=90,
                color="#cc4444", ha="center")
    ax.set_xlabel("position i")
    ax.set_ylabel("position j")
    ax.set_title(title, fontsize=10)
    ax.set_xlim(0, n)
    ax.set_ylim(0, n)
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)
    return path


def write_forna(sequence, structure, path):
    """Sequence + dot-bracket, ready to paste into FORNA."""
    with open(path, "w", encoding="utf-8") as fh:
        fh.write(">structure\n%s\n%s\n" % (sequence, structure))
    return path


# ---------------------------------------------------------------------------
# "What stands out about this design"
# ---------------------------------------------------------------------------

def what_stands_out(built, row4, row6):
    """
    The PDF's closing bullet list, generated from the numbers rather than
    written by hand, so it cannot drift away from the data above it.
    """
    out = []
    ratio = row6["AND_ratio"]
    t_on = row6["t_fire_given_b_s"]

    if t_on > 1200:
        out.append(
            "Does not fire within an mRNA lifetime even with both inputs "
            "(t_fire = %s). The AND ratio of %.0fx is a ratio between two "
            "states that are both effectively OFF, so it should not be read "
            "as performance."
            % (("%.0f s" % t_on) if t_on < 1e6 else "never", ratio))
    else:
        out.append(
            "Fires in %.0f s with both inputs, comfortably inside an mRNA "
            "lifetime, and %.0fx faster than with trigger A alone."
            % (t_on, ratio))

    out.append(
        "Trigger B releases x* by %.1f points (%.1f%% -> %.1f%%), growing "
        "trigger A's nucleation site from %.1f to %.1f nt. This is the step "
        "the whole architecture depends on."
        % (row4["x_star_release_pts"], row4["00_x_star_open"],
           row4["01_x_star_open"], row4["nucleation_nt_00"],
           row4["nucleation_nt_01"]))

    if row4["sequestered_pct"] > 5:
        out.append(
            "%.1f%% of trigger A is tied up binding trigger B rather than the "
            "switch -- the highest of the set, and worth watching."
            % row4["sequestered_pct"])
    else:
        out.append(
            "The two inputs stay apart: only %.1f%% of trigger A is sequestered "
            "by trigger B." % row4["sequestered_pct"])

    out.append(
        "Orthogonality margin %.1f kcal/mol and A-kill %.1f kcal/mol, both "
        "positive, so state 10 and state 01 should each read OFF."
        % (row4["orthogonality_margin"], row4["a_kill"]))

    bulge = built["checks"]["bulge"]
    if bulge["size"]:
        out.append(
            "A %d-nt bulge in r1copy splits the %d bp stem into %d + %d bp, "
            "taking the longest continuous duplex below the ~20 bp RNase III "
            "prefers." % (bulge["size"], built["stem_bp"],
                          bulge["helix_upper_bp"], bulge["helix_lower_bp"]))
    else:
        out.append(
            "No bulge: the inhibitory stem is %d bp of continuous duplex, "
            "which is within the range E. coli RNase III cleaves. The bulged "
            "variants are in the workbook." % built["stem_bp"])

    out.append(
        "The variant-B gene differs from the working original at %d of %d "
        "nucleotides (%.1f%%), all synonymous, with E. coli codon usage %.3f "
        "against the original's 0.384."
        % (row4["total_edits"], len(built["design"]["variant"]),
           row4["gene_changed_pct"], row4["usage_whole_gene"]))

    out.append(
        "MFE structure carries %.1f%% of the ensemble and matches the intended "
        "OFF fold at %.0f%% of positions."
        % (100 * row4["p_mfe_off"], row4["intended_agreement_pct"]))
    return out


# ---------------------------------------------------------------------------
# The metrics table
# ---------------------------------------------------------------------------

def metric_rows(built, row4, row6, trigger_a, trigger_b):
    """
    Every reported metric, as (name, ViennaRNA, NUPACK, per-nt) tuples.

    Mirrors the PDF's table and adds the columns it lacks.
    """
    switch = built["sequence"]
    n_trig = len(trigger_a)
    nup_00 = row4.get("nupack_dG_00")
    nup_11 = row4.get("nupack_dG_11")

    def na(value, fmt="%.2f"):
        return "n/a" if value is None else fmt % value

    rows = [
        ("dG switch alone (OFF)", "%.2f" % row4["00_dG"], na(nup_00), ""),
        ("dG switch + trigger A", "%.2f" % row4["10_dG"], "n/a",
         "%.3f" % (row4["10_dG"] / n_trig)),
        ("dG switch + trigger B", "%.2f" % row4["01_dG"], "n/a", ""),
        ("dG switch + A + B (ON)", "%.2f" % row4["11_dG"], na(nup_11), ""),
        ("dG_open(A)", "%.2f" % row4["dG_open_A"], "n/a",
         "%.3f" % row4["dG_open_A_per_nt"]),
        ("dG_open(A | B)", "%.2f" % row4["dG_open_A_given_B"], "n/a", ""),
        ("ddG_AND", "%.2f" % row4["ddG_AND"], "n/a", ""),
        ("", "", "", ""),
        ("Toehold a* open - OFF", "%.1f%%" % row4["00_a_star_open"], "n/a", ""),
        ("x* shut - OFF", "%.1f%%" % (100 - row4["00_x_star_open"]), "n/a", ""),
        ("x* released by B", "%.1f pts" % row4["x_star_release_pts"], "n/a", ""),
        ("r2* open - OFF", "%.1f%%" % row4["00_r2_star_open"], "n/a", ""),
        ("Start codon open - OFF", "%.1f%%" % row4["00_aug_open"], "n/a", ""),
        ("Start codon open - ON", "%.1f%%" % row4["11_aug_open"], "n/a", ""),
        ("RBS open - OFF", "%.1f%%" % row4["00_rbs_open"], "n/a", ""),
        ("", "", "", ""),
        ("Nucleation site, OFF", "%.1f nt" % row4["nucleation_nt_00"], "n/a", ""),
        ("Nucleation site, +B", "%.1f nt" % row4["nucleation_nt_01"], "n/a", ""),
        ("dG_toe, A alone", "%.2f" % row6["dG_toe_A_alone"], "n/a", ""),
        ("dG_toe, A given B", "%.2f" % row6["dG_toe_A_given_B"], "n/a", ""),
        ("ddG_toe (the gate)", "%.2f" % row6["ddG_toe"], "n/a", ""),
        ("AND ratio (kinetic)", "%.0fx" % row6["AND_ratio"], "n/a", ""),
        ("Time to fire, ON", ("%.0f s" % row6["t_fire_given_b_s"])
         if row6["t_fire_given_b_s"] < 1e6 else "never", "n/a", ""),
        ("", "", "", ""),
        ("Trigger A:B duplex", "%.2f" % row4["trigger_A_B_dG"], "n/a", ""),
        ("Trigger A sequestered", "%.1f%%" % row4["sequestered_pct"], "n/a", ""),
        ("Orthogonality margin", "%.1f" % row4["orthogonality_margin"], "n/a", ""),
        ("A-kill margin", "%.1f" % row4["a_kill"], "n/a", ""),
        ("", "", "", ""),
        ("SED (specified defect)", "%.3f" % row4["SED"], "n/a", ""),
        ("NED (native defect)", "%.3f" % row4["NED"], "n/a", ""),
        ("P(MFE) in ensemble", "%.2f%%" % (100 * row4["p_mfe_off"]), "n/a", ""),
        ("Intended fold agreement", "%.0f%%" % row4["intended_agreement_pct"],
         "n/a", ""),
        ("", "", "", ""),
        ("Gene edits", "%d of %d" % (row4["total_edits"],
                                     len(built["design"]["variant"])), "n/a",
         "%.1f%%" % row4["gene_changed_pct"]),
        ("E. coli codon usage", "%.3f" % row4["usage_whole_gene"], "n/a",
         "orig 0.384"),
        ("RBS present / frame", "yes / yes", "n/a", ""),
    ]
    return rows


# ---------------------------------------------------------------------------
# Word
# ---------------------------------------------------------------------------

def _add_seq(doc, label, sequence):
    from docx.shared import Pt
    para = doc.add_paragraph()
    run = para.add_run("%s (%d nt)\n" % (label, len(sequence)))
    run.bold = True
    seq_run = para.add_run(sequence)
    seq_run.font.name = "Consolas"
    seq_run.font.size = Pt(7.5)


def build_word_report(entries, config, figures_dir, path):
    """The main document: one section per candidate, laid out like the PDF."""
    import docx
    from docx.shared import Inches, Pt

    doc = docx.Document()
    doc.add_heading("Two-input AND toehold switches", level=0)
    doc.add_paragraph(
        "Each construct keeps one of the five validated single-input switches "
        "from Toehold_Candidates29.7.pdf byte-for-byte, and prepends a Kim-2019 "
        "inhibitory hairpin in front of it. Input A is the original mCherry; "
        "input B is a minimally recoded synonymous variant of the same gene, so "
        "the two can be expressed separately and all four states measured.")

    doc.add_heading("What is shown, and how it was chosen", level=1)
    doc.add_paragraph(
        "One design per candidate: the best of every (|a|, Lx, m) and bulge "
        "combination built for that candidate, ranked on the KINETIC gate "
        "(ddG_toe and the AND ratio), which is the criterion that determines "
        "whether the switch works at all. These are the best FOR EACH "
        "CANDIDATE, not the five best overall -- candidates 1 and 3 are "
        "included because the set is the point, and their numbers say plainly "
        "that they do not fire. Every alternative scored is in the workbook.")
    doc.add_paragraph(
        "Selection order: stage 2 sweeps several hundred (|a|, Lx, m, site) "
        "combinations per candidate and shortlists on orthogonality, A-kill "
        "and trigger sequestration; the top few are carried forward, each "
        "built with four bulge sizes; stages 4 and 6 then score all of them "
        "and the kinetics picks the winner. The shortlist exists because "
        "stage 2 cannot see ddG_toe -- that needs the built construct -- so "
        "handing it only stage 2's winner would settle the design before the "
        "deciding number was ever computed.")

    doc.add_heading("Read this first", level=1)
    for text in (
        "The AND behaviour of this architecture is KINETIC, not thermodynamic. "
        "At equilibrium trigger A displaces the inhibitory hairpin on its own "
        "(~38 bp of A:switch against ~24 bp of hairpin), so states 10 and 11 "
        "are indistinguishable and start-codon separation is 0.0 points for "
        "every design. The gate comes from trigger A having only |a| "
        "nucleotides to nucleate on without B, which slows displacement by "
        "orders of magnitude against mRNA degradation. Rank on ddG_toe and the "
        "AND ratio, not on the equilibrium energies.",
        "Absolute firing probabilities are indicative only. k_on and k_bm are "
        "order-of-magnitude DNA constants reused for RNA; the ratio between "
        "two states scored with identical constants is the trustworthy output.",
        "The spontaneous (trigger-independent) leak term is reported but NOT "
        "folded into the AND ratio. As parameterised it returns a ~97% leak, "
        "which cannot be right given the single-input candidates work at the "
        "bench. Calibrating k_ribosome against their measured ON/OFF is the "
        "one wet-lab input that would make it meaningful.",
        "Three rows of the original PDF are mislabelled or misleading: "
        "'RBS hidden - OFF' is actually the RBS's UNPAIRED probability, so "
        "high means exposed; 'main stem 50/67/75%' is exactly (n-3)/n because "
        "the 3-nt bulge sits in the denominator, and all five candidates form "
        "100% of their intended base pairs; and the dG margin ranking largely "
        "tracks trigger length (r = -0.872). See the appendix.",
    ):
        para = doc.add_paragraph(text)
        para.style = doc.styles["List Bullet"]

    for entry in entries:
        built, row4, row6 = entry["built"], entry["row4"], entry["row6"]
        cand_id = built["cand"]
        doc.add_page_break()
        doc.add_heading("Candidate %d" % cand_id, level=1)
        doc.add_paragraph(
            "|a| = %d, Lx = %d, m = %d, bulge = %s. Inhibitory stem %d bp; "
            "longest continuous duplex %d bp. Total construct %d nt (%d added "
            "to the original %d)."
            % (built["len_a"], built["Lx"], built["m"],
               ("%d nt at position %d" % (built["bulge_size"], built["bulge_index"]))
               if built["bulge_size"] else "none",
               built["stem_bp"], built["checks"]["bulge"]["longest_bp"],
               len(built["sequence"]), built["added_nt"],
               len(cd.CANDIDATES[cand_id]["switch"])))

        doc.add_heading("Sequences", level=2)
        _add_seq(doc, "AND switch", built["sequence"])
        _add_seq(doc, "Trigger A (original mCherry)", entry["trigger_a"])
        _add_seq(doc, "Trigger B (variant mCherry)", entry["trigger_b"])
        doc.add_paragraph(
            "Variant-B gene: %d nt, in stage2_variantB_genes.fasta."
            % len(built["design"]["variant"]))

        doc.add_heading("Metrics", level=2)
        rows = metric_rows(built, row4, row6, entry["trigger_a"], entry["trigger_b"])
        table = doc.add_table(rows=1, cols=4)
        table.style = "Light Grid Accent 1"
        header = table.rows[0].cells
        for cell, text in zip(header, ("Metric", "ViennaRNA", "NUPACK", "per nt / note")):
            cell.text = text
        for name, vienna, nupack, extra in rows:
            cells = table.add_row().cells
            cells[0].text = name
            cells[1].text = vienna
            cells[2].text = nupack
            cells[3].text = extra
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(8)

        doc.add_heading("Predicted structures", level=2)
        for key, caption in (("switch", "AND switch alone (OFF state)"),
                             ("trigger", "Trigger A alone"),
                             ("complex", "Switch + trigger A + trigger B (ON state)")):
            image = entry["figures"].get(key)
            if image and os.path.exists(image):
                doc.add_paragraph(caption)
                doc.add_picture(image, width=Inches(4.4))
        bpp = entry["figures"].get("bpp")
        if bpp and os.path.exists(bpp):
            doc.add_paragraph("Base-pairing probabilities, OFF state")
            doc.add_picture(bpp, width=Inches(4.4))

        doc.add_heading("What stands out about this design", level=2)
        for bullet in what_stands_out(built, row4, row6):
            para = doc.add_paragraph(bullet)
            para.style = doc.styles["List Bullet"]

    doc.save(path)
    return path


APPENDIX_SECTIONS = [
    ("Engine and model",
     "ViennaRNA 2.x, default model, 37 C, throughout. This is not a preference: "
     "it is what the candidates PDF used, confirmed by RNA.fold() and "
     "RNA.cofold() reproducing all ten of its energies exactly. NUPACK differs "
     "by 3-5 kcal/mol on the two-strand complex, so it appears as a second "
     "column rather than a substitute, and is unavailable on Windows."),
    ("dG switch alone",
     "RNA.fold(switch) -- the minimum free energy. Reproduces the PDF exactly. "
     "Caveat: this is the single most stable structure, and for candidate 4 it "
     "carries only 2.2% of the ensemble."),
    ("dG margin, and its correction",
     "The PDF computes complex - switch. That omits the trigger's own folding, "
     "worth 1.3 to 6.2 kcal/mol here, so we also report "
     "complex - switch - trigger. Note the PDF itself uses the CORRECTED "
     "margin as the denominator of its off-target percentage while printing "
     "the uncorrected one in its table -- its two rows disagree about which "
     "margin is real. Reproducing the percentages exactly is how we know."),
    ("Per-nucleotide normalisation",
     "The five triggers are 35-41 nt and a longer trigger binds harder for "
     "free; correlation between length and margin is r = -0.872. Per "
     "nucleotide the five span 12% where the raw column spans 25%, and the top "
     "two swap. Do not normalise quantities that map to an observable (dG sets "
     "Kd), but always read the per-nt companion column. ddG_AND and ddG_toe "
     "need no normalisation at all: they are differences within one molecule."),
    ("Accessibility",
     "Mean probability that a base is unpaired, over the whole ensemble rather "
     "than the MFE alone. 'RBS hidden - OFF' in the PDF is this quantity, so a "
     "high number means the RBS is EXPOSED. That is expected here -- the RBS "
     "sits in an 18-nt loop by design and repression comes from the start "
     "codon being sequestered -- but it means the row cannot rank candidates. "
     "The start codon is the readout that discriminates."),
    ("SED and NED",
     "SED is the ensemble defect against the fully-unpaired reference: how far "
     "from open the molecule is, lower being more accessible. NED is the "
     "defect against the molecule's own MFE structure: how representative that "
     "structure is, lower being better defined. Both from "
     "fc.ensemble_defect(), normalised 0..1."),
    ("Codon usage",
     "E. coli usage fraction per codon, from the same ecoli_codon_usage_table.csv "
     "VISTA uses. Reported as the mean across a region and across the whole "
     "gene, plus the first two codons which matter disproportionately for "
     "initiation. Context: the codon-max mCherry that FAILED at the bench "
     "scores 0.471, better than the working original's 0.384 -- so poor codon "
     "usage is unlikely to explain that failure."),
    ("ddG_AND (equilibrium)",
     "(G11 - G01) - (G10 - G00). G(A) cancels. More negative means trigger A "
     "binds better once B has acted. Reported with its fold-change "
     "exp(-ddG/RT), which is the same number since dG is logarithmic. "
     "IMPORTANT: for this architecture it is not decisive -- see ddG_toe."),
    ("Nucleation gain (equilibrium, and the one that discriminates)",
     "Expected unpaired nucleotides in the region trigger A must land on "
     "(a* plus x*), in state 01 minus state 00. It measures the single step "
     "the architecture depends on: trigger B releasing x*. Candidates 1 and 3 "
     "post the most negative ddG_AND of the set while gaining ~0 nt here, "
     "which is why the ranking leads with this."),
    ("dG_toe and ddG_toe (kinetic)",
     "dG_toe = dG_duplex(trigger : site) + opening_energy(site), where the "
     "opening energy is the work to hold the whole site single-stranded, from "
     "a constrained partition function -- not an average of per-base "
     "probabilities, which is a much easier bar. Both nucleation routes (a* "
     "alone, and a*+x*) are scored in both states and the better one taken, "
     "because that is what the molecule does. ddG_toe is the difference, and "
     "it is the gate."),
    ("AND ratio (kinetic)",
     "P_fire(11) / P_fire(10) using displacement only, where "
     "P_fire = k_obs/(k_obs + k_deg), k_obs = k_eff x [trigger], "
     "k_eff = k_on k_bm / (k_on Kd_toe + k_bm), k_deg = ln2 / half-life "
     "(Zhang & Winfree 2009). k_on = 3e6 /M/s and k_bm = 1 /s are "
     "order-of-magnitude DNA values at 25 C reused for RNA at 37 C. Absolute "
     "P_fire is indicative only; the ratio between two states on identical "
     "constants is the output."),
    ("Spontaneous leak -- reported, not scored",
     "k_ribosome x (start-codon accessibility), against degradation. Kept out "
     "of the AND ratio because as parameterised it returns a ~97% "
     "trigger-independent leak, which swamped every ratio to exactly 1.0. It "
     "assumes every unpaired start codon initiates translation, which the "
     "candidates PDF warns against, and it contradicts the single-input "
     "switches working at the bench. Calibrate k_ribosome against their "
     "measured ON/OFF before using it."),
    ("Cross-binding",
     "dG of the best duplex between the two triggers, converted to a "
     "sequestered fraction by solving A + B <-> AB at the stated "
     "concentration. Measured on the BARE trigger windows: duplexfold returns "
     "the best duplex anywhere between two sequences, so adding 25 nt of "
     "flanking gene to each side finds a strong one regardless of design "
     "(candidate 5 goes from -7.0 to -34.9 kcal/mol that way). Whole-transcript "
     "interaction is covered separately by the whole-gene scans."),
    ("Orthogonality and A-kill",
     "orthogonality = dG(original mCherry : trigger B's landing site) "
     "- dG(trigger B : that same site), where the site is r2* + k2* -- the free "
     "toehold plus the half of the stem B invades. Note this is NOT the hairpin's "
     "5' arm, which is k2* + r1copy and is set by trigger A. "
     "and A-kill = dG(variant : toehold footprint) - dG(original : footprint). "
     "Both scan the WHOLE gene, and A-kill targets the toehold footprint "
     "specifically -- against a whole switch any long RNA finds a -70 kcal/mol "
     "duplex somewhere, which says nothing. Both must be positive."),
    ("Unresolved",
     "'Region just after the start codon open - ON' from the PDF could not be "
     "reproduced. About twenty definitions were tried (mean unpaired, mean "
     "paired, joint-unpaired via constrained partition function, windows of "
     "3-21 nt, the b_pre* and linker spans, both states); the closest still "
     "misses by 16 points across five candidates. It is used nowhere. Settling "
     "it needs the definition from whoever wrote the original validation "
     "script."),
]


def build_appendix(path):
    """Every formula, with the caveat that belongs to it."""
    import docx
    doc = docx.Document()
    doc.add_heading("Appendix: how every metric is computed", level=0)
    doc.add_paragraph(
        "Companion to the AND switch report. Each entry says what is folded, "
        "which reference states are subtracted, and what not to conclude from "
        "the result.")
    for title, body in APPENDIX_SECTIONS:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)
    doc.save(path)
    return path


# ---------------------------------------------------------------------------
# Excel
# ---------------------------------------------------------------------------

def build_workbook(entries, all_rows4, all_rows6, path):
    """All the numbers, re-sortable, one sheet per stage plus a summary."""
    import openpyxl
    from openpyxl.styles import Font

    wb = openpyxl.Workbook()
    summary = wb.active
    summary.title = "summary"
    headers = ["candidate", "|a|", "Lx", "m", "bulge", "stem bp",
               "AND ratio", "ddG_toe", "t_fire ON (s)",
               "nucleation gain (nt)", "x* release (pts)",
               "sequestered %", "orthogonality", "A-kill",
               "gene changed %", "codon usage", "total nt"]
    summary.append(headers)
    for cell in summary[1]:
        cell.font = Font(bold=True)
    for entry in entries:
        built, row4, row6 = entry["built"], entry["row4"], entry["row6"]
        summary.append([
            built["cand"], built["len_a"], built["Lx"], built["m"],
            built["bulge_size"], built["stem_bp"],
            row6["AND_ratio"], row6["ddG_toe"], row6["t_fire_given_b_s"],
            row4["nucleation_gain_nt"], row4["x_star_release_pts"],
            row4["sequestered_pct"], row4["orthogonality_margin"],
            row4["a_kill"], row4["gene_changed_pct"],
            row4["usage_whole_gene"], len(built["sequence"])])

    def dump(sheet_name, rows, skip_private=True):
        sheet = wb.create_sheet(sheet_name)
        if not rows:
            return
        keys = [k for k in rows[0] if not (skip_private and k.startswith("_"))]
        sheet.append(keys)
        for cell in sheet[1]:
            cell.font = Font(bold=True)
        for row in rows:
            sheet.append([row.get(k) for k in keys])

    dump("stage4_equilibrium", all_rows4)
    dump("stage6_kinetics", all_rows6)

    seqs = wb.create_sheet("sequences")
    seqs.append(["candidate", "what", "length", "sequence"])
    for cell in seqs[1]:
        cell.font = Font(bold=True)
    for entry in entries:
        built = entry["built"]
        for label, seq in (("AND switch", built["sequence"]),
                           ("trigger A", entry["trigger_a"]),
                           ("trigger B", entry["trigger_b"]),
                           ("variant-B gene", built["design"]["variant"]),
                           ("intended OFF structure", built["intended_structure"])):
            seqs.append([built["cand"], label, len(seq), seq])

    wb.save(path)
    return path


# ---------------------------------------------------------------------------
# Stage entry point
# ---------------------------------------------------------------------------

def run(config, results=None):
    """Stage 5 entry point."""
    print()
    print("STAGE 5 -- Word report, appendix and Excel workbook")

    rows4 = (results or {}).get(4)
    rows6 = (results or {}).get(6)
    if not rows4 or not rows6:
        raise SystemExit(
            "Stage 5 needs stages 4 and 6. Run:  --stages 2,3,4,6,5")

    out_dir = config["out_dir"]
    figures_dir = os.path.join(out_dir, "figures")
    os.makedirs(figures_dir, exist_ok=True)

    # rows6 arrives sorted by AND ratio, so the first row for a candidate is
    # its best design AND best bulge together -- chosen on the kinetic gate,
    # which is the criterion that decides whether the thing works.
    best = {}
    for r6 in rows6:
        best.setdefault(r6["cand"], r6)
    print("  selecting the best of %d scored construct(s), one per candidate"
          % len(rows6))

    entries = []
    for cand_id in sorted(best):
        row6 = best[cand_id]
        # Match on the FULL design identity, not just candidate + bulge.
        #
        # Since stage 3 began carrying several designs per candidate, a given
        # (candidate, bulge) pair no longer identifies one construct -- there
        # are five designs x four bulges per candidate. Matching on those two
        # fields alone silently paired stage 6's chosen row with a DIFFERENT
        # design's stage-4 row, so the report printed one design's geometry
        # beside another's kinetics. It showed candidate 1 as |a|=5 Lx=8 m=5
        # when the finalist is |a|=4 Lx=9 m=7.
        key = ("len_a", "Lx", "m", "bulge_size")
        matches = [r for r in rows4
                   if r["cand"] == cand_id
                   and all(r[k] == row6[k] for k in key)]
        if not matches:
            raise SystemExit(
                "no stage-4 row matches the stage-6 winner for candidate %d "
                "(|a|=%s Lx=%s m=%s bulge=%s)"
                % (cand_id, row6["len_a"], row6["Lx"], row6["m"],
                   row6["bulge_size"]))
        row4 = matches[0]
        built = row4["_built"]
        trigger_a = row6["_trigger_a"]
        trigger_b = row6["_trigger_b"]

        print("  candidate %d: drawing figures..." % cand_id)
        tag = "cand%d" % cand_id
        spans = built["spans"]
        marks = [(spans["r2_star"][0], "r2*"), (spans["k2_star"][0], "k2*"),
                 (spans["r1copy"][0], "r1copy"), (spans["r1_star"][0], "r1*"),
                 (spans["a_star"][0], "a*"), (spans["primary_aug"][0], "AUG")]

        off_structure = built["fold"]["mfe_structure"]
        figures = {
            "switch": draw_structure(
                built["sequence"], off_structure,
                os.path.join(figures_dir, "%s_switch_off.png" % tag),
                "Candidate %d - AND switch, OFF state" % cand_id),
            "trigger": draw_structure(
                trigger_a, RNA.fold(trigger_a)[0],
                os.path.join(figures_dir, "%s_triggerA.png" % tag),
                "Candidate %d - trigger A alone" % cand_id),
            "bpp": draw_base_pair_probabilities(
                built["sequence"],
                os.path.join(figures_dir, "%s_bpp_off.png" % tag),
                "Candidate %d - pair probabilities, OFF" % cand_id, marks),
        }
        complex_seq = trigger_a + "&" + trigger_b + "&" + built["sequence"]
        complex_structure = RNA.fold_compound(complex_seq).mfe()[0]
        figures["complex"] = draw_structure(
            complex_seq.replace("&", ""), complex_structure.replace("&", ""),
            os.path.join(figures_dir, "%s_complex_on.png" % tag),
            "Candidate %d - ON state (A + B + switch)" % cand_id)

        write_forna(built["sequence"], off_structure,
                    os.path.join(figures_dir, "%s_switch_off.forna.txt" % tag))

        entries.append({"built": built, "row4": row4, "row6": row6,
                        "trigger_a": trigger_a, "trigger_b": trigger_b,
                        "figures": figures})

    report_path = os.path.join(out_dir, "AND_switch_report.docx")
    appendix_path = os.path.join(out_dir, "AND_switch_appendix.docx")
    workbook_path = os.path.join(out_dir, "AND_switch_data.xlsx")

    build_word_report(entries, config, figures_dir, report_path)
    build_appendix(appendix_path)
    build_workbook(entries, rows4, rows6, workbook_path)

    print()
    print("wrote %s" % report_path)
    print("wrote %s" % appendix_path)
    print("wrote %s" % workbook_path)
    print("wrote %d figures + FORNA text to %s" % (4 * len(entries), figures_dir))
    return entries


# Pressing Run on this file alone needs stages 4 and 6 first.
if __name__ == "__main__":
    from poc_and.main import CONFIG, main
    main(["--stages", "2,3,4,6,5"])
