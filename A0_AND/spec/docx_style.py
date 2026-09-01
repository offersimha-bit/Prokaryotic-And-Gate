"""Shared styling helpers for the A0_AND documents.

Both build_spec.py (the full specification) and build_brief.py (the one-page-per-
topic brief) import from here, so the two documents always look the same.
Nothing here knows anything about the content.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
FIG = os.path.join(ROOT, "figures")

CONTENT_CM = 17.0
NAVY = RGBColor(0x1B, 0x3A, 0x5E)
GREY = RGBColor(0x49, 0x50, 0x57)
RED = RGBColor(0xA4, 0x13, 0x3C)
HDR_FILL = "1B3A5E"
ALT_FILL = "F1F4F8"
NOTE_FILL = "FFF6E5"
DEC_FILL = "EAF3EC"


def new_document(margin_cm=2.0):
    """An A4 document with the house Normal style already applied."""
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"
    st.font.size = Pt(10)
    st.paragraph_format.space_after = Pt(6)
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(margin_cm)
    sec.top_margin = sec.bottom_margin = Cm(margin_cm)
    return doc


def save(doc, path):
    """Save, falling back to path_1, path_2 ... when Word or OneDrive holds the file."""
    base, ext = os.path.splitext(path)
    for attempt in range(6):
        target = path if attempt == 0 else "%s_%d%s" % (base, attempt, ext)
        try:
            doc.save(target)
            if target != path:
                print("NOTE: %s was locked (open in Word?); wrote %s instead"
                      % (os.path.basename(path), os.path.basename(target)))
            return target
        except PermissionError:
            continue
    raise PermissionError("could not write %s or any fallback name" % path)


# ----------------------------------------------------------------- helpers
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def cell_text(cell, text, bold=False, size=8.5, color=None, mono=False,
              align=None):
    """Writes text into a cell. Honours \\n as a line break and ** as bold."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if align is not None:
        p.alignment = align
    for i, chunk in enumerate(str(text).split("\n")):
        if i:
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            if align is not None:
                p.alignment = align
        for seg, seg_bold in parse_bold(chunk):
            r = p.add_run(seg)
            r.bold = bold or seg_bold
            r.font.size = Pt(size)
            r.font.name = "Consolas" if mono else "Calibri"
            if color is not None:
                r.font.color.rgb = color


def repeat_header(t):
    """mark row 0 as a header row so it repeats when the table splits a page"""
    trPr = t.rows[0]._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def table(doc, headers, rows, widths, size=8.5, mono_cols=(), zebra=True,
          row_fills=None, tail=True):
    """widths in cm; must sum to CONTENT_CM.

    tail=False omits the spacer paragraph that normally follows a table. Use it
    for the last table before a page break: an empty Normal paragraph is half a
    centimetre tall and is enough on its own to spill onto a blank page.
    """
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for j, (h, w) in enumerate(zip(headers, widths)):
        c = t.rows[0].cells[j]
        c.width = Cm(w)
        shade(c, HDR_FILL)
        cell_text(c, h, bold=True, size=size, color=RGBColor(0xFF, 0xFF, 0xFF))
    if any(headers):
        repeat_header(t)
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j, (v, w) in enumerate(zip(row, widths)):
            cells[j].width = Cm(w)
            fill = None
            if row_fills and row_fills[i]:
                fill = row_fills[i]
            elif zebra and i % 2 == 1:
                fill = ALT_FILL
            if fill:
                shade(cells[j], fill)
            cell_text(cells[j], v, size=size, mono=(j in mono_cols))
    if tail:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


def h1(doc, text, new_page=False):
    p = doc.add_paragraph()
    if new_page:
        # a break carried on the heading itself, rather than an extra empty
        # paragraph, which can otherwise spill and produce a blank page
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = NAVY
    pPr = p._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "1B3A5E")
    bdr.append(bottom)
    pPr.append(bdr)


def h2(doc, text, new_page=False):
    p = doc.add_paragraph()
    if new_page:
        p.paragraph_format.page_break_before = True
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    return p


def para(doc, text, size=10, italic=False, color=None, space_after=6,
         bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    for seg, is_bold in parse_bold(text):
        r = p.add_run(seg)
        r.bold = is_bold or bold
        r.italic = italic
        r.font.size = Pt(size)
        if color is not None:
            r.font.color.rgb = color
    return p


def parse_bold(text):
    """split on ** markers into (text, bold) segments"""
    out, buf, bold = [], "", False
    i = 0
    while i < len(text):
        if text[i:i + 2] == "**":
            if buf:
                out.append((buf, bold))
            buf, bold, i = "", not bold, i + 2
        else:
            buf += text[i]
            i += 1
    if buf:
        out.append((buf, bold))
    return out


def bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.7)
    for seg, is_bold in parse_bold(text):
        r = p.add_run(seg)
        r.bold = is_bold
        r.font.size = Pt(size)
    return p


def mono(doc, text, size=9, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def callout(doc, title, text, fill=NOTE_FILL):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    t.autofit = False
    c = t.rows[0].cells[0]
    c.width = Cm(CONTENT_CM)
    shade(c, fill)
    c.text = ""
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(9.5)
    p2 = c.add_paragraph()
    p2.paragraph_format.space_after = Pt(3)
    for seg, is_bold in parse_bold(text):
        r = p2.add_run(seg)
        r.bold = is_bold
        r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def figure(doc, filename, caption, width_cm=CONTENT_CM):
    doc.add_picture(os.path.join(FIG, filename), width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    for seg, is_bold in parse_bold(caption):
        r = p.add_run(seg)
        r.bold = is_bold
        r.italic = not is_bold
        r.font.size = Pt(8.5)
        r.font.color.rgb = GREY
