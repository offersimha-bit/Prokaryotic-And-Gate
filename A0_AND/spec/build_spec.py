"""Build A0_AND_Specification.docx.

This script is the source of truth for the specification document.
Edit the text here and re-run; do not hand-edit the .docx.

    python spec/build_spec.py

Figures are read from ../figures/ and are produced by make_fig_architecture.py,
make_fig_overlap.py and make_fig_loops.py.  Styling lives in docx_style.py.
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_style import (new_document, save, table, h1, h2, para, bullet, mono,
                        callout, figure, shade, cell_text,
                        CONTENT_CM, NAVY, GREY, RED, HDR_FILL, ALT_FILL,
                        NOTE_FILL, DEC_FILL, FIG, ROOT)

OUT = os.path.join(ROOT, "A0_AND_Specification.docx")

# ----------------------------------------------------------------- document
doc = new_document()

# ---- title block
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("A0 AND Gate")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(14)
r = p.add_run("Architecture specification — for review and approval")
r.font.size = Pt(13)
r.font.color.rgb = GREY

table(doc, ["", ""], [
    ["Document", "Complete definition of a two-input AND riboregulator: domains, "
                 "lengths, design rules, and the calculations that will rank candidates."],
    ["Status", "Draft for supervisor review. Section 9 lists twelve open decisions."],
    ["Prerequisites", "None. The document is self-contained."],
    ["Date", "1 September 2026"],
], [3.2, 13.8], size=9.5, zebra=True)

callout(doc, "What is being asked",
        "Everything in sections 1–8 is fixed and ready to implement. "
        "**Section 9 lists twelve decisions we need from you** before code is written and "
        "constructs are ordered. Each carries our recommendation and the reason it matters. "
        "Six are ready to sign off as recommended; six are genuine choices where we do not "
        "think the evidence settles it.")

# ================================================================= 1
h1(doc, "1.  What we are building")
para(doc, "A single mRNA that produces GFP only when **two different RNA inputs are "
          "present at the same time**. One input is a window of gene A, the other a window "
          "of gene B. Neither input alone should produce GFP.")
para(doc, "The construct is a **toehold switch** carrying a second, **inhibitory hairpin** "
          "immediately upstream of it. The main hairpin sequesters the ribosome binding site "
          "and the start codon. The inhibitory hairpin sequesters the site where the main "
          "trigger would otherwise land.")
bullet(doc, "**Trigger B** binds a free single-stranded toehold and opens the inhibitory "
            "hairpin. This exposes the landing site — and nothing else.")
bullet(doc, "**Trigger A** then binds that newly exposed site and opens the main hairpin, "
            "releasing the start codon.")
para(doc, "The name **A0** refers to the one parameter that defines this variant: the "
          "unpaired spacing "
          "between the two hairpins, conventionally called a, is set to **zero**. "
          "Section 2 explains why that specific number.")

# ================================================================= 2
h1(doc, "2.  The two architectures this is built on")
para(doc, "The geometry below was **measured, not quoted**: each published construct was "
          "rebuilt from its supplementary sequences and folded with ViennaRNA 2.7.2 at 37 °C. "
          "This matters because the published text does not state the stem lengths directly, "
          "and they are the numbers we need.")

h2(doc, "2.1  Measured geometry")
table(doc,
      ["", "Green — Series A\n(tsgen2 / VISTA)", "Kim — main\nhairpin (G3n*)",
       "Kim — inhibitory\nhairpin (G5)"],
      [["stem below the internal loop", "9 bp", "11 bp", "11 bp"],
       ["internal loop", "3 × 3", "1 × 1", "1 × 1"],
       ["stem above the internal loop", "6 bp", "6 bp", "8 bp"],
       ["total base pairs", "15 bp", "17 bp", "19 bp"],
       ["**arm span, bulge included**", "**18 nt**", "**18 nt**", "20 nt"],
       ["start codon", "fully unpaired,\na 3-nt bulge", "only 1 nt unpaired;\nthe other 2 "
                                                        "are paired", "— (no RBS or AUG)"],
       ["hairpin loop", "11 nt (carries RBS)", "15 nt (carries RBS)", "15 nt (no RBS)"],
       ["trigger length", "36 nt", "30 nt", "30 nt"],
       ["free single-stranded toehold", "30 nt", "15 nt", "13 nt"],
       ["invasion into the stem", "6 bp", "15 bp", "17 bp"]],
      [5.6, 4.0, 3.7, 3.7], size=9,
      row_fills=[None, None, None, None, "FFF0CC", None, None, None, None, None])
para(doc, "Table 1 — the three published hairpins. The highlighted row is the quantity this "
          "specification fixes (§4.1). Note the last row: **Kim's start codon is not a "
          "3-nt bulge** — read against their own structure specification, only its first "
          "nucleotide is unpaired and the other two are base-paired inside the lower stem. "
          "Green's is fully unpaired, and ours follows Green.", size=8.5, italic=True,
     color=GREY)

h2(doc, "2.2  Kim's key result, and why a = 0")
para(doc, "Kim held trigger A at 30 nt and slid its binding site upstream, trading two "
          "quantities against each other: **a**, the unpaired spacing between the two "
          "hairpins, and *****, the number of nucleotides of trigger A's site that lie "
          "buried inside the inhibitory hairpin's stem. In their constructs a + * = 15. "
          "Three variants were tested:")
table(doc, ["construct", "a", "*", "behaviour reported"],
      [["Sw-G5-G3n5", "10", "5", "trigger B only modulates the output range — "
                                 "effectively still a one-input switch"],
       ["Sw-G5-G3n8", "7", "8", "intermediate"],
       ["Sw-G5-G3n11", "4", "11", "\"a primary input molecule of the two-input AND gate\""]],
      [3.4, 1.3, 1.3, 11.0], size=9,
      row_fills=[None, None, DEC_FILL])
para(doc, "The trend is monotonic and it is the entire rationale for A0: **the smaller a is, "
          "the more the construct behaves as a genuine AND gate**. A0 takes a to its limit. "
          "With a = 0, trigger A has no exposed nucleotide anywhere on the switch until "
          "trigger B has acted — the dependency becomes structural rather than a matter of degree.")
callout(doc, "The one caveat on this",
        "a = 0 is one step beyond Kim's shortest tested spacing. The direction is supported "
        "by their data; the specific value is an extrapolation. §5.3 therefore includes an "
        "a = 4 construct as a benchmark, so the panel contains a point that is known to work.")

# ================================================================= 3
h1(doc, "3.  The A0 architecture")
figure(doc, "fig1_architecture.png",
       "**Figure 1.** The A0 switch and its two triggers. Domains that pair with each "
       "other share a colour. Both arms span 18 nt; the main arm is 9 + 3-nt bulge + 6. "
       "**|x| = 9 is drawn only as an example** — len_x is the one free design parameter and "
       "len_k2 = 18 − len_x moves with it. §4 gives the general rule.", width_cm=14.2)

h2(doc, "3.1  The switch, 5' → 3'")
para(doc, "Every element of the construct, in order. \"Pairs with\" describes the OFF state.")
table(doc,
      ["#", "element", "pairs with", "length (nt)", "sequence determined by"],
      [["1", "cap", "— free", "3", "fixed: GGG"],
       ["2", "r2*", "— free (trigger-B toehold)", "len_r2", "revcomp(r2); window of gene B"],
       ["3", "Secondary_pre*", "x*  (#9)", "len_x", "= x; window of gene A"],
       ["4", "k2*", "SecondaryZ  (#8)", "18 − len_x", "revcomp(k2); window of gene B"],
       ["5", "SUSA", "SUSD  (#7)", "0", "free design — placeholder"],
       ["6", "Secondary loop", "— free", "15", "fixed; must contain no RBS-like motif"],
       ["7", "SUSD", "SUSA  (#5)", "0", "revcomp(SUSA) — placeholder"],
       ["8", "SecondaryZ", "k2*  (#4)", "18 − len_x", "≈ k2, deliberately weakened (§6.3)"],
       ["9", "x*", "Secondary_pre*  (#3)", "len_x", "revcomp(x); window of gene A"],
       ["10", "a", "— free", "0", "spacer — placeholder, fixed at 0 in A0"],
       ["11", "Main_pre*", "Main_pre  (#19)", "len_main_pre = **9**",
        "revcomp(Main_pre); gene A"],
       ["12", "opposing_bulge", "— free in OFF;\npaired by trigger A in ON", "3",
        "revcomp of trigger A's 3 nt (§6.4)"],
       ["13", "k1*", "MainZ  (#17)", "len_k1 = **6**",
        "revcomp(k1); window of gene A"],
       ["14", "MUSA", "MUSD  (#16)", "0", "free design — placeholder"],
       ["15", "RBS loop", "— free", "18", "fixed"],
       ["16", "MUSD", "MUSA  (#14)", "0", "revcomp(MUSA) — placeholder"],
       ["17", "MainZ", "k1*  (#13)", "len_k1 = **6**",
        "≈ k1, deliberately weakened (§6.3)"],
       ["18", "AUG", "— free", "3", "fixed: start codon"],
       ["19", "Main_pre", "Main_pre*  (#11)", "len_main_pre = **9**",
        "window of gene A"],
       ["20", "LINKER", "— free", "21", "fixed"],
       ["21", "GFP CDS", "— free", "759", "fixed"]],
      [0.8, 3.2, 4.3, 2.7, 6.0], size=8)
para(doc, "Table 2 — the 21 elements of the switch, at the recommended geometry "
          "**len_main_pre = 9, len_k1 = 6** (§4.1 and §6.6; decisions D1 and D2, both still "
          "open). **Only len_x and len_r2 remain free**, and len_k2 = 18 − len_x moves with "
          "len_x. Elements 5, 7, 10, 14 and 16 have length 0: they are **placeholders**, "
          "kept in the layout so that a future version can add an upper stem to either "
          "hairpin, or re-open the inter-hairpin spacer, without renumbering anything.",
     size=8.5, italic=True, color=GREY)

h2(doc, "3.2  The triggers")
mono(doc, "Trigger A  (main, gene A)       5'-  k1 - opposing_bulge - Main_pre - a - x  -3'")
mono(doc, "Trigger B  (secondary, gene B)  5'-  k2 - Secondary_pre - r2               -3'")
para(doc, "**Reading rule.** A trigger binds antiparallel, so its 5' end pairs with the "
          "**3'-most** element of its footprint on the switch. Listing a trigger's domains "
          "in switch order gives the reverse of the correct sequence. This rule is why "
          "trigger A begins with k1 even though k1* sits deep inside the main hairpin.")
para(doc, "**The shared domain.** Secondary_pre = x*. Trigger A carries x; trigger B carries "
          "x*. The two triggers are therefore complementary to each other over len_x "
          "nucleotides. That overlap is what couples the two halves of the gate, and it is "
          "also the single largest risk in the design (§5.1).")

h2(doc, "3.3  Mechanism")
figure(doc, "fig2_truthtable.png",
       "**Figure 2.** The four input states. Only 11 produces GFP.")
para(doc, "In state 11 trigger A's entire nucleation budget is len_x nucleotides — that is "
          "all trigger B exposes. Once nucleated, trigger A branch-migrates through "
          "Main_pre* and k1*, displacing Main_pre and MainZ and freeing the start codon.")

# ================================================================= 4
h1(doc, "4.  Lengths")

h2(doc, "4.1  Decision: both hairpin arms span 18 nt")
para(doc, "**Both hairpin arms are fixed at 18 nucleotides, counting the bulge** — not 18 "
          "base pairs. The distinction decides how deep the hairpin is, so the measurement "
          "is worth stating plainly.")
para(doc, "The two published main hairpins are built differently: Green's is 9 bp + a 3-nt "
          "bulge + 6 bp, Kim's is 11 bp + a 1-nt bulge + 6 bp. Their base-pair counts differ "
          "(15 against 17) but **both arms span exactly 18 nucleotides**. So 18 is not a "
          "compromise between two published numbers — it is the number both of them share.")
mono(doc, "Green   9 bp  +  3-nt bulge  +  6 bp   =  18 nt per arm,  15 bp paired\n"
          "Kim    11 bp  +  1-nt bulge  +  6 bp   =  18 nt per arm,  17 bp paired\n"
          "A0      9 bp  +  3-nt bulge  +  6 bp   =  18 nt per arm,  15 bp paired")
para(doc, "A0 takes Green's arrangement, which is also the one all five of our own tested "
          "switches use. Reading 18 as a base-pair count instead would give a 21-nt arm — "
          "deeper than anything either group, or we, have ever built. This is decision "
          "**D1**.")
para(doc, "The secondary hairpin has no bulge, so for it the two readings coincide: 18 nt "
          "of arm, 18 bp paired.")

h2(doc, "4.2  What 18 bp determines")
para(doc, "Fixing the stem closes every remaining length in the design. Each stem is split "
          "in two by an unpaired region, and the two halves must sum to 18:")
mono(doc, "main hairpin       len_main_pre  +  3 (bulge)  +  len_k1  =  18\n"
          "secondary hairpin  len_x         +  len_k2                =  18")
para(doc, "So len_k1 and len_k2 are not free parameters at all — they are consequences:")
table(doc, ["quantity", "rule", "worked values"],
      [["len_k1", "6, fixed by the SD→AUG spacing (§6.6)", "→ len_main_pre = 18 − 3 − 6 = 9"],
       ["len_k2", "18 − len_x", "|x| = 4 → 14   |x| = 9 → 9   |x| = 14 → 4"],
       ["length of trigger A", "len_k1 + 3 + len_main_pre + len_x  =  18 + len_x",
        "|x| = 4 → 22 nt   |x| = 12 → 30 nt   |x| = 14 → 32 nt"],
       ["length of trigger B", "18 + len_r2", "len_r2 = 32 → 50 nt"],
       ["len_main_pre", "multiple of 3 (§6.1)", "9"]],
      [3.6, 5.0, 8.4], size=9)
para(doc, "Two of these are worth pausing on. **Trigger A's length is 18 + len_x regardless "
          "of how the main arm is split** — the split moves nucleotides between k1 and "
          "Main_pre without changing the total. And at |x| = 12 trigger A is exactly 30 nt, "
          "the length both Green and Kim used. **|x| is the one free design parameter**; "
          "every length quoted elsewhere in this document at |x| = 9 is an illustration, not "
          "a fixed value.")
para(doc, "**len_k1 is not free either.** §6.6 shows that k1 is exactly the distance between "
          "the Shine–Dalgarno sequence and the start codon, which is a determinant of "
          "translation initiation rate. Green, Kim and all five of our own tested switches "
          "put that distance at 6 nt. Setting len_k1 = 6 then forces len_main_pre = 9, and "
          "every length in the construct except len_x and len_r2 is closed.")

h2(doc, "4.3  Trigger B — the split is now determined")
para(doc, "Trigger B's ~50 nt could in principle be split many ways between k2, "
          "Secondary_pre and r2, with very different behaviour. **The arm span of §4.1 "
          "closes it**, and the reasoning is worth stating so the gap is not "
          "re-introduced.")
para(doc, "Trigger B is k2 + Secondary_pre + r2. Two of those three are already pinned: "
          "Secondary_pre = x* has length len_x, and k2 has length 18 − len_x. Their sum is "
          "18 whatever len_x turns out to be. So:")
mono(doc, "len_r2  =  (length of trigger B)  −  18")
para(doc, "At the target length of 50 nt, len_r2 = 32. Before the stem was fixed, both k2 "
          "and r2 floated, and a 50-nt trigger could have been split in many ways with very "
          "different behaviour — that was the actual gap, not a missing number.")
callout(doc, "One number here does deserve a second look",
        "len_r2 = 32 nt is trigger B's free toehold. Kim used **13 nt** for the same role; "
        "Green used 30 nt for a toehold that had no hairpin above it. A 32-nt toehold is "
        "long, which buys specificity but also makes trigger B's binding much stronger than "
        "trigger A's. This is decision **D3**.")

# ================================================================= 5
h1(doc, "5.  The overlap, and what goes to the lab")

h2(doc, "5.1  The overlap x is the central parameter")
para(doc, "len_x is simultaneously two things, pulling in opposite directions:")
bullet(doc, "It is trigger A's **only** nucleation site in state 11 — so longer is better.")
bullet(doc, "It is the region over which the two triggers are **complementary to each other** "
            "— so longer is worse, because they bind each other instead of the switch.")
figure(doc, "fig3_overlap.png",
       "**Figure 3.** Fraction of trigger A tied up in a trigger-A : trigger-B duplex, "
       "against overlap length. Red: perfect complementarity — unusable beyond |x| = 8. "
       "Green: the same overlaps carrying evenly spread mismatches, which is what moves the "
       "curve back down. The green curve is jagged and turns up again past |x| = 14 because "
       "a fixed proportion of mismatches does not hold the longest complementary run down "
       "as the overlap grows; §5.1.1 states the rule actually used.")
para(doc, "With perfect complementarity the useful range ends abruptly: sequestration goes "
          "from 0.6 % at |x| = 7 to 19 % at 8 and 70 % at 9. **Introducing mismatches moves "
          "the whole curve**, and beyond |x| ≈ 8 they are the only thing that keeps the "
          "design usable.")
para(doc, "**Placement matters more than the ratio.** Enumerating all 220 ways to place three "
          "mismatches in a 12-nt overlap, over 120 random sequences each, the same "
          "mismatch count spans a 9.5 kcal/mol range depending only on where the mismatches "
          "sit. What tracks that range is the **longest uninterrupted complementary run**:")
table(doc, ["longest complementary run inside a 12-nt overlap",
            "3 nt", "4 nt", "5 nt", "6 nt", "7 nt", "8 nt", "9 nt"],
      [["trigger A sequestered (10 nM, 37 °C)",
        "0.0 %", "0.1 %", "1.1 %", "4.0 %", "19 %", "61 %", "90 %"]],
      [5.6, 1.63, 1.63, 1.63, 1.63, 1.63, 1.63, 1.62], size=8.5, zebra=False)
callout(doc, "Why run length is a diagnostic and not a threshold",
        "A single 7-nt complementary run gives 0.6 % sequestration. But a 15-nt overlap with "
        "**one** mismatch in the middle also has a longest run of 7 — two of them — and "
        "measures **100 %**, because the two short helices stitch together across the "
        "mismatch at almost no cost. Run length tracks sequestration well within a fixed "
        "overlap length (r = −0.79, and the count of matched positions r = −0.85) but does "
        "not bound it across lengths. It is reported next to every design to explain the "
        "ranking; the number that decides is the measured duplex energy itself.")
h2(doc, "5.1.1  The mismatch rule")
para(doc, "Mismatches inside the overlap are what make the search viable, so they are "
          "allowed — but not without shape. The rule is:")
table(doc, ["|x|", "mismatches allowed", "everywhere subject to"],
      [["4 – 5", "**none**", "—"],
       ["6 – 10", "**up to 2**", "never at either end of the overlap;\n"
                                  "never more than 2 consecutive"],
       ["11 – 14", "**up to 3**", "the same two conditions"]],
      [2.6, 4.4, 10.0], size=9)
para(doc, "**Why the two placement conditions.** A mismatch at either end costs nothing "
          "to the trigger-trigger duplex — the folding routine simply trims the frayed end "
          "— so it would be a mismatch that buys no decoupling while still weakening trigger "
          "B's grip and fraying the junction where the secondary hairpin meets its free "
          "toehold. Three or more consecutive mismatches open an internal loop large enough "
          "to stall trigger B's branch migration through Secondary_pre. Both conditions "
          "therefore keep the mismatches useful rather than merely tolerated.")
para(doc, "**What the rule buys.** Counting every (window of gene A, window of gene B) pair "
          "between our two real transcripts, with and without it:")
table(doc, ["|x|", "perfectly complementary sites", "sites under the rule", "widening"],
      [["4 – 5", "1,716 and 431", "unchanged — no mismatches allowed", "—"],
       ["6", "95", "7,062", "74×"],
       ["7", "24", "2,737", "114×"],
       ["8", "3", "1,009", "336×"],
       ["9 – 14", "**0 at every length**", "359, 113, 244, 87, 26, 8",
        "from nothing to something"]],
      [2.2, 5.4, 5.4, 4.0], size=9,
      row_fills=[None, None, None, None, "EAF3EC"])
para(doc, "The last row is the important one: **for this gene pair there is not a single "
          "perfectly complementary site at any overlap of 9 nt or more.** Without mismatches "
          "the entire long-overlap half of the design range is empty, and the sweep over "
          "|x| that §5.3 proposes could not be run at all.")
callout(doc, "What the rule does not do",
        "It does not bound sequestration, and it is not meant to. Enumerating every "
        "placement it permits: the median design stays under 2 % out to |x| = 12, but the "
        "worst permitted placement reaches 93 % at |x| = 10 and the median is 45 % by "
        "|x| = 14 — because clustering the allowed mismatches to one side still leaves one "
        "long clean run. Sequestration is therefore measured and **ranked** (§7.4), not "
        "filtered. Kim's own working trigger pair sits at 73 %, so a filter would be the "
        "wrong instrument in any case.")
para(doc, "**One thing worth keeping in mind: we do not choose where the mismatches fall.** "
          "Both triggers are windows of natural transcripts, so the mismatch pattern is "
          "whatever the two genes happen to differ by. The rule decides which "
          "(window, window) pairs count as a usable overlap; it does not design anything. "
          "That is also why the spread in the callout above is a statement about which gene "
          "pairs will rank badly, not about a choice we could get wrong.")

h2(doc, "5.2  The search over gene pairs")
para(doc, "Trigger A and trigger B are windows of two chosen input RNAs. Not every pair of "
          "genes admits a design: the two windows must be complementary over at least 4 nt, "
          "because that overlap is the mechanism. The search is therefore:")
bullet(doc, "For every window of gene A and every window of gene B, find all overlaps of "
            "length **|x| = 4 to 14**, admitting mismatches under the rule of §5.1.1.")
bullet(doc, "**Every value of |x| from 4 upwards is designed and scored — not only the "
            "longest available overlap.** The longest overlap is often the worst choice, "
            "because it is also the most sequestered.")
bullet(doc, "Each surviving overlap determines the whole construct through §4.2.")
para(doc, "**Mismatches earn their place from |x| = 6 upwards, not only at long overlaps.** "
          "They buy nothing thermodynamically at that length — a 6-nt perfect overlap is "
          "already at 0 % sequestration — and each one costs trigger B about 4.5 kcal/mol of "
          "grip on the switch. They are allowed because they are what makes the search find "
          "anything at all. Counting every (window of A, window of B) pair between our two "
          "real transcripts, with no rule applied:")
table(doc, ["overlap |x|", "perfect", "≤ 1 mismatch", "≤ 2 mismatches", "≤ 3 mismatches"],
      [["6 nt", "95", "2,014", "16,324", "77,853"],
       ["7 nt", "24", "539", "5,552", "30,884"],
       ["8 nt", "3", "148", "1,781", "11,586"],
       ["9 nt", "0", "34", "574", "4,154"],
       ["12 nt", "0", "0", "8", "160"],
       ["14 nt", "0", "0", "0", "12"]],
      [3.4, 3.4, 3.4, 3.4, 3.4], size=9,
      row_fills=[None, None, None, "FFF6E5", None, None])
para(doc, "Table 3 — candidate overlap sites between GFP (759 nt) and mCherry (711 nt), two "
          "unrelated genes. **At |x| = 9 there is not one perfectly complementary site in the "
          "whole pair; allowing a single mismatch produces 34.** Even at |x| = 6, permitting "
          "one mismatch widens the pool 21-fold — which matters because most of those sites "
          "will then be lost to the frame rule, the in-frame-stop filter and the structure "
          "checks. Restricting mismatches to long overlaps would leave the search with "
          "almost nothing to rank.", size=8.5, italic=True, color=GREY)

h2(doc, "5.3  What to sweep in the constructs sent to the lab")
para(doc, "Constructs are expensive, so the sweep should spend them only where the outcome "
          "is genuinely uncertain. We measured which parameters actually move the "
          "thermodynamics:")
table(doc, ["parameter", "range", "sweep?", "why"],
      [["len_x  (overlap)", "4 – 14", "YES — 4 levels",
         "The dominant axis, and the same axis Kim swept. Trades nucleation against "
         "sequestration; the optimum is not predictable."],
       ["ddG_pref  (§6.3)", "0 or ≈ 2 kcal/mol", "YES — 2 levels",
        "Weakening MainZ helps the ON state and hurts the OFF state by the same amount. "
        "A genuine knife-edge."],
       ["len_main_pre / len_k1 split", "9 + 3 + 6, fixed", "NO — the arm span fixes it",
        "§4.1 sets the arm span at 18 nt and §6.6 sets len_k1 = 6 from the SD→AUG spacing, "
        "which leaves len_main_pre = 9. Thermodynamically the split is inert in any case "
        "(−9.61 against −9.57 kcal/mol of drive), so it would not be worth a construct even "
        "if it were free."],
       ["a  (inter-hairpin spacer)", "0", "NO — but see below",
        "Fixed at 0 by definition of A0."],
       ["Secondary loop, SUSA/MUSA", "fixed", "NO", "Held constant so the sweep stays interpretable."]],
      [3.7, 2.8, 2.8, 7.7], size=8.5)
para(doc, "**Proposed panel — 12 constructs.**")
table(doc, ["group", "n", "detail"],
      [["Gate designs", "8", "4 levels of |x| (4, 7, 10, 13) × 2 levels of ddG_pref (0, 2)"],
       ["Kim benchmark", "1", "the best |x| repeated at a = 4, matching Kim's working AND "
                              "construct — a positive control for the architecture itself"],
       ["Control: no inhibitory hairpin", "1", "main hairpin alone — a plain one-input "
                                               "toehold switch. Gives the ON ceiling."],
       ["Control: dead trigger-B site", "1", "r2* scrambled so trigger B cannot bind. "
                                             "Isolates leak through the main hairpin."],
       ["Control: no hairpins", "1", "RBS + AUG + linker + GFP. Maximum-expression reference."]],
      [5.4, 1.2, 10.4], size=9)
para(doc, "Each of the 8 gate designs is measured in all four input states, so the panel is "
          "12 constructs and 39 measurements. Note that **each gate design needs its own "
          "gene-B sequence** if the input genes are being recoded to create the overlap; if "
          "the gene pairs are natural, they do not.")

# ================================================================= 6
h1(doc, "6.  Design rules")

h2(doc, "6.1  Reading frame")
para(doc, "Translation starts at the AUG (#18) and runs AUG | Main_pre | LINKER | GFP.")
para(doc, "**Rule: the total length of every element between the start codon and the GFP CDS "
          "must be divisible by 3.** Today that is Main_pre + LINKER. LINKER is 21 nt, which "
          "is already a multiple of 3, so the rule reduces to len_main_pre being a multiple "
          "of 3 — but it is stated in the general form so it still holds if an element is "
          "ever inserted there.")

h2(doc, "6.2  No in-frame stop codon in Main_pre")
para(doc, "Main_pre is a literal stretch of gene A, and it is translated as an N-terminal "
          "extension of GFP. Its codons are dictated by whichever trigger window is chosen, "
          "so a window carrying an in-frame stop would kill the reporter.")
para(doc, "**Rule: reject any trigger-A window whose Main_pre introduces an in-frame stop "
          "codon.** This is a hard filter inside the search, not a check afterwards. Green "
          "applied the same filter when tiling their library, and Kim wrote it into their "
          "NUPACK script as a pattern constraint — then had to disable it, because the "
          "designer has no knowledge of the reading frame. It has to live in our own code.")
para(doc, "The cost is modest. At len_main_pre = 9, measured directly on the real "
          "transcripts, the filter removes **10.1 % of candidate windows in GFP and 6.5 % in "
          "mCherry** — both below the 13.4 % a random sequence of that length would give, "
          "because coding sequence is depleted of stop codons in every frame.")

h2(doc, "6.3  MainZ and SecondaryZ must lose to the trigger")
para(doc, "MainZ holds k1* shut in the OFF state, and trigger A has to displace it. "
          "SecondaryZ does the same for k2* against trigger B. Both are copies of the "
          "trigger's own domain, which raises a subtle problem.")
para(doc, "A **perfectly Watson–Crick** copy binds its partner slightly harder than the real "
          "trigger does, because the real trigger sequence generally contains at least one "
          "G·U wobble. The trigger would then be climbing uphill to displace the very clamp "
          "meant to release it. Copying the trigger's own sequence exactly makes the exchange "
          "isoenergetic — neither favoured nor disfavoured.")
para(doc, "**The rule is one-sided.** MainZ and SecondaryZ must never bind their partner "
          "*more tightly* than the real trigger does. Two settings satisfy that and both are "
          "on the table:")
bullet(doc, "**ddG_pref = 0** — Z is an exact copy of the trigger's own domain, wobbles "
            "included, so the exchange is isoenergetic and neither side is favoured. This is "
            "the simplest construction and is a legitimate choice, not a fallback.")
bullet(doc, "**ddG_pref > 0** — Z is weakened further by a G·U wobble, a single mismatch or "
            "a 1-nt bulge, so the trigger is actively preferred.")
para(doc, "What is excluded is only the third case: a **perfectly Watson–Crick** Z, which "
          "would bind harder than the real trigger and make the exchange uphill. ddG_pref is "
          "swept over both allowed settings (§5.3) rather than fixed.")
callout(doc, "Trade-off to measure, not assume",
        "Every weakening of Z also weakens the OFF-state lock by the same amount. ddG_pref "
        "is therefore swept rather than fixed (§5.3), and the OFF-state leak must be reported "
        "next to the ON-state gain for each level.")

h2(doc, "6.4  The opposing bulge is the ON driving force")
para(doc, "The 3-nt opposing_bulge (#12) and the AUG (#18) face each other across the main "
          "stem, forming a 3 × 3 internal loop. In the OFF state those six nucleotides are "
          "unpaired. In the ON state trigger A pairs straight through them, forming 18 "
          "contiguous base pairs where the hairpin managed only 15 plus an internal loop.")
para(doc, "This is not a side effect — it is what makes the gate thermodynamically possible "
          "at all, given that MainZ is otherwise a copy of trigger A's own domain. Measured "
          "over 400 random sequences, at a fixed 18-nt arm span:")
table(doc, ["main hairpin geometry", "hairpin ΔG", "trigger A : 5' arm ΔG",
            "advantage to trigger A"],
      [["15 bp, no bulge (control)", "−23.23", "−25.15", "−1.92 kcal/mol"],
       ["9 + 3-nt bulge + 6  (Green, and A0)", "−21.94", "−31.55", "−9.61 kcal/mol"],
       ["11 + 1-nt bulge + 6  (Kim)", "−24.99", "−31.44", "−6.45 kcal/mol"],
       ["12 + 3-nt bulge + 6  (a deeper arm, for scale)", "−28.46", "−38.03",
        "−9.57 kcal/mol"]],
      [5.6, 3.2, 4.4, 3.8], size=9,
      row_fills=[None, DEC_FILL, None, None])
para(doc, "Three things follow. Without a bulge the gate has a 1.9 kcal/mol margin — too "
          "thin to rely on. A 3-nt bulge is worth about 3 kcal/mol more than Kim's 1-nt "
          "bulge, which is why we keep Green's arrangement. And the drive is insensitive to "
          "how deep the arm is (−9.61 at 9 + 3 + 6 against −9.57 at 12 + 3 + 6) — but the "
          "**OFF lock is not** (−21.94 against −28.46). A shallower arm is the looser lock, "
          "and the four-state table will show that directly.", space_after=10)

h2(doc, "6.5  GFP's own start codon")
para(doc, "The GFP CDS begins with its own ATG. Translation initiates upstream at the "
          "switch's AUG, so in the fusion that ATG becomes an internal methionine. Keep it "
          "or delete it? Our three sources disagree, which is why this is a decision and not "
          "a rule. Both published architectures keep the reporter's own ATG directly after "
          "the 21-nt linker, in frame:")
mono(doc, "Kim    ... AACCTGGCGGCAGCGCAAAAG | ATGAGTAAAGGAGAAGAACTTTTCACTGG ...\n"
          "Green  ... AACCTGGCGGCAGCGCAAAAG | ATGCGTAAAGGAGAAGAACTTTTCACTGG ...")
para(doc, "**Our own lab's switch generator does the opposite.** The script that produced "
          "our five tested switches passes the reporter CDS in as sequence[3:], commented "
          "\"remove the AUG\" — so the constructs we have bench data for have no reporter "
          "start codon, only the switch's own.")
para(doc, "**Neither option is thermodynamically consequential.** Folding each of the five "
          "tested switches with 90 nt of GFP attached, with and without the reporter ATG, "
          "moves the joint unpaired probability of the switch's own start codon by at most "
          "0.09, and in four of five cases by under 0.02.")
para(doc, "**Two worries can be dismissed outright**, and it is worth recording that they "
          "were checked rather than assumed:")
bullet(doc, "**A second ribosome entry site at GFP's ATG.** There is no Shine–Dalgarno-like "
            "motif anywhere in the 21-nt linker, and the switch's own SD sits 33 nt upstream, "
            "far outside initiation range. Keeping GFP's ATG does not create a second start.")
bullet(doc, "**Cryptic initiation further inside GFP.** Our GFP does carry SD-like motifs "
            "(AAGGAG at position 8, AGGAGA at position 9), but scanning all 759 nt finds "
            "**no AUG at initiation spacing — 3 to 12 nt — downstream of any of them**. "
            "There is no internal start site, with or without the leading ATG.")
para(doc, "What is left is a preference. **We recommend deleting it**: it matches the "
          "generator behind the switches we have data for, and it leaves the construct with "
          "exactly one start codon, which makes the accessibility metric of §7.2 "
          "unambiguous. The frame is unaffected either way (759 and 756 are both divisible "
          "by 3). This is decision **D4**.")


h2(doc, "6.6  The RBS loop")
para(doc, "Both reference designs build this loop the same way: a **fixed 11-nt ribosome "
          "binding site sitting flush against the loop's 3' end**, preceded by a free flank "
          "that the NUPACK designer fills in. Our own generator does exactly this too — the "
          "loop is declared as N(L−11) followed by the constant AACAGAGGAGA. The three "
          "loops turn out to be strictly nested:")
figure(doc, "fig4_loops.png",
       "**Figure 4.** Top: the RBS loop in the three designs, aligned on their 3' ends. "
       "Bottom: because the RBS is 3'-flush, the Shine–Dalgarno-to-start-codon spacing is "
       "exactly len_k1.", width_cm=16.4)
table(doc, ["design", "loop", "length", "designed flank", "fixed RBS"],
      [["Green 2014 / VISTA", "AACAGAGGAGA", "11 nt", "— none", "AACAGAGGAGA"],
       ["Kim 2019", "CAAGAACAGAGGAGA", "15 nt", "CAAG", "AACAGAGGAGA"],
       ["ours (from the supervisor)", "AGACAAGAACAGAGGAGA", "18 nt", "AGACAAG",
        "AACAGAGGAGA"]],
      [4.4, 4.6, 1.8, 2.6, 3.6], size=8.5, mono_cols=(1, 3, 4))
para(doc, "**Our loop is Kim's loop with AGA added, and Kim's is Green's with CAAG added.** "
          "There is no methodological disagreement between the three to resolve — only a "
          "choice of how much designed flank to put in front of a constant RBS.")
para(doc, "**First consequence: the RBS being 3'-flush is what fixes len_k1.** Nothing "
          "separates the Shine–Dalgarno sequence from the start codon except k1 itself, so "
          "the spacing that governs translation initiation *is* len_k1. Measured directly:")
table(doc, ["construct", "SD → AUG spacing"],
      [["Green / VISTA fixed scaffold", "6 nt"],
       ["Kim Sw-G5-G3n11", "6 nt"],
       ["our five tested switches (candidates 1–5)", "6 nt in all five"],
       ["**A0 with len_k1 = 6** (→ len_main_pre = 9)", "**6 nt**"],
       ["A0 with len_k1 = 9 (→ len_main_pre = 6)", "9 nt — no precedent in this scaffold"]],
      [10.0, 7.0], size=9,
      row_fills=[None, None, None, "EAF3EC", "FFF6E5"])
para(doc, "This is what fixes **len_k1 = 6**: it is the only value that reproduces the "
          "spacing used by both published designs and by every switch we have tested at the "
          "bench, and the split is thermodynamically inert either way (§6.4), so nothing is "
          "paid for it. Together with the 18-nt arm span of §4.1 it leaves "
          "**len_main_pre = 18 − 3 − 6 = 9**, which is also a multiple of 3 as §6.1 requires.")
para(doc, "**Second consequence: loops are not portable.** We substituted each of the three "
          "loops into all five of our tested switches, changing nothing else, and folded the "
          "result:")
table(doc, ["loop substituted in", "mean P(RBS entirely unpaired)",
            "intended hairpin still formed?"],
      [["Green, 11 nt", "0.77", "yes, in 5 of 5"],
       ["Kim, 15 nt", "0.33", "**no — in 0 of 5**"],
       ["ours, 18 nt", "0.68", "yes, in 5 of 5"]],
      [4.6, 6.4, 6.0], size=9,
      row_fills=[None, "FFE0E0", "EAF3EC"])
para(doc, "Kim's loop does not merely score worse — dropped into our sequence context it "
          "refolds the whole switch into a different global structure, in every one of the "
          "five. A loop that is correct in one construct can be wrong in the next. Whatever "
          "loop is chosen, **the intended fold has to be re-verified for every design** "
          "(assertion 2, §8).")
callout(doc, "The one mild surprise, for you to weigh",
        "Green's shorter 11-nt loop gives slightly **better** RBS accessibility than ours in "
        "four of the five tested switches (mean 0.77 vs 0.68). The margin is small and our "
        "18-nt loop is the one with bench data behind it, so **our recommendation is to keep "
        "the 18-nt loop unchanged** — but you should know the shorter loop measures "
        "marginally better before signing that off. This is decision **D12**.")
para(doc, "**Third consequence: the pattern-prevention list needs a scope.** Kim forbade "
          "aaaa, cccc, gggg, uuuu and the six-long IUPAC runs kkkkkk, mmmmmm, rrrrrr, ssssss, "
          "wwwwww, yyyyyy; our generator uses the identical list. But the RBS itself is "
          "AGGAGA — six purines — so it violates rrrrrr. **The list must be applied to "
          "designed positions only, never to fixed domains**, or the designer will refuse to "
          "place the ribosome binding site.")

h2(doc, "6.7  The inhibitory loop")
para(doc, "Kim used two different methods for this loop in the same paper, and it matters "
          "which one we are copying.")
table(doc, ["Kim construct series", "inhibitory loop", "method"],
      [["Sw-G5-G3n* (Kim's Figure 1) — **the series this architecture is based on**",
        "CAAGAACUUAGACAA, 15 nt, fixed",
        "carried over from an earlier Green-derived switch; not designed for this paper"],
       ["Sw-T0-T3 (Kim's Figure 2) — a different construct", "designed de novo",
        "NUPACK, declared as domain a = N21 with the prevent list above"]],
      [5.6, 4.6, 6.8], size=8.5)
para(doc, "Two details of the de-novo route are easy to misread, so they are recorded here "
          "exactly as the supporting information gives them. First, **N21 is not the loop**: "
          "read against the structure string, those 21 free nucleotides span 3 nt of the "
          "hairpin's upper stem, the 12-nt loop, and 6 nt of the closing strand. Second, that "
          "hairpin has a 12-bp stem and a 12-nt loop, not the 19-bp stem and 15-nt loop of "
          "Kim's Figure 1 series — the two are different constructs.")
para(doc, "**One constraint, not two.** The loop must carry no Shine–Dalgarno-like motif, "
          "because that would create a second ribosome entry point. It does **not** need to "
          "be free of AUG codons: the secondary loop sits in the 5'UTR, bacterial initiation "
          "requires an SD at the right spacing, and an AUG with no SD upstream of it "
          "initiates nothing, so no constraint is placed on start codons in this loop.")
para(doc, "**Recommendation.** Start from Kim's fixed 15-nt loop, which we verified carries "
          "no SD-like motif, and check the intended fold for each design. "
          "Where it fails that check, design the loop de novo in NUPACK exactly as Kim did — "
          "free N nucleotides, the prevent list applied to designed positions only, and the "
          "intended structure as the design target. §6.6 shows this failure is not "
          "hypothetical: a fixed loop that suits one sequence context can break the next, and "
          "the secondary hairpin's stem is set by the gene pair and therefore changes with "
          "every design. This is decision **D7**.")

# ================================================================= 7
h1(doc, "7.  How designs will be evaluated")

h2(doc, "7.1  Report the centroid, not only the MFE")
para(doc, "For each of the four states we report the MFE structure and its free energy, "
          "**the probability of the MFE structure**, the **centroid** structure with its "
          "energy and its distance from the MFE, the ensemble free energy, and the ensemble "
          "diversity.")
para(doc, "The reason is that P(MFE) collapses with length, and our constructs are long. "
          "Measured on a real two-hairpin switch, extended stepwise:")
table(doc, ["what is folded", "length", "P(MFE)", "centroid distance from MFE"],
      [["one hairpin alone", "70 nt", "0.44", "0.9 bp"],
       ["two-hairpin switch, no reporter", "129 nt", "0.023", "6 bp"],
       ["switch + linker + first 90 nt of GFP", "240 nt", "1.4 × 10⁻⁴", "28 bp"],
       ["full construct with GFP", "909 nt", "9.5 × 10⁻¹⁵", "151 bp"]],
      [6.6, 2.4, 3.2, 4.8], size=9,
      row_fills=[None, None, None, "FFE0E0"])
para(doc, "**At full length the MFE structure is one microstate in 10¹⁴ and describes "
          "essentially nothing.** Any ranking that reads a single structure at that length is "
          "reading noise. The operative descriptions are the centroid and the base-pairing "
          "probability matrix; the MFE remains meaningful only for the isolated hairpin "
          "regions, where P(MFE) is still order 0.1–0.5.")
para(doc, "The intended OFF structure is written directly from the element spans in Table 2, "
          "independently of any folding, and compared against **both** the MFE and the "
          "centroid. Both agreements are reported.")
para(doc, "**This is not a theoretical concern — the two descriptions already disagree on the "
          "only question we actually ask.** Folding our five tested switches at three lengths "
          "and asking each structure whether the start codon is paired:")
table(doc, ["what is folded", "P(MFE)", "MFE and centroid agree on 'is the AUG paired?'"],
      [["switch alone (95–107 nt)", "0.02 – 0.08", "5 of 5"],
       ["switch + 90 nt of GFP (185–197 nt)", "10⁻⁴ – 10⁻²", "5 of 5"],
       ["full construct with GFP (854–866 nt)", "≈ 3 × 10⁻¹⁵", "**3 of 5**"]],
      [6.4, 3.6, 7.0], size=9,
      row_fills=[None, None, "FFE0E0"])
para(doc, "Where P(MFE) is order 10⁻² the two agree everywhere. Where it is order 10⁻¹⁵ they "
          "disagree about whether the gate is open in two of five constructs — and there is "
          "no basis in the numbers for preferring the MFE's answer. **Which structure the "
          "ranking reads is decision D11**, and we recommend the centroid together with the "
          "base-pairing probability matrix at full length, with the MFE reported but used "
          "only over the isolated hairpin region where it still means something.")

h2(doc, "7.2  Accessibility around the start codon")
para(doc, "**Recommendation: report the probability that a defined window around the start "
          "codon is entirely unpaired at once**, not the average over its bases.")
mono(doc, "P_open(W)   =  Z( all of W unpaired )  /  Z\n"
          "dG_open(W)  =  - RT · ln P_open(W)")
para(doc, "**Why this rather than the usual metric.** The common measure averages the "
          "unpaired probability across the window's bases. A ribosome does not sample bases "
          "independently — it needs its whole footprint free at the same moment. For a window "
          "whose bases are each 90 % unpaired, the averaged metric always reads 0.90, while "
          "the joint probability lies anywhere between 0.90 (if the bases open together) and "
          "far less (if they open independently). The two can rank designs differently, and "
          "only the joint probability corresponds to the physical requirement.")
para(doc, "**Windows**, all anchored on the start codon and reported as a nested series:")
bullet(doc, "the AUG alone (3 nt);")
bullet(doc, "the RBS through the AUG — the Shine–Dalgarno spacing region;")
bullet(doc, "AUG through AUG + 3n, for n = 1 … 10 codons;")
bullet(doc, "the standard 30-nt ribosome footprint.")
para(doc, "The rankable quantity is the difference between states: "
          "ΔΔG_open = dG_open(state 11) − dG_open(worst of 00, 01, 10). Reporting the whole "
          "nested series rather than one window shows how far the opening extends, which is "
          "what distinguishes a switch that merely exposes the AUG from one that clears the "
          "ribosome's path.")
para(doc, "**Implementation.** ViennaRNA computes this directly: build a fold compound, apply "
          "hard constraints forcing the window unpaired, and take the free-energy difference "
          "between the constrained and unconstrained ensembles. One extra partition function "
          "per window.")

h2(doc, "7.3  Four-state energetics")
para(doc, "The truth table is evaluated as a multi-strand equilibrium over the four states, "
          "reporting for each: the ensemble free energy, the base-pairing probability matrix, "
          "and the accessibility series of §7.2. The headline discrimination is")
mono(doc, "ddG_AND  =  [ dG_open(11) - dG_open(01) ]  -  [ dG_open(10) - dG_open(00) ]")
para(doc, "which measures how much more trigger A opens the switch when trigger B is present "
          "than when it is absent. Length cancels in the difference, so this number needs no "
          "normalisation.")

h2(doc, "7.4  Trigger–trigger sequestration, and the loops as accidental trigger sites")
para(doc, "Reported at a stated concentration (10 nM per transcript unless changed), as both "
          "the duplex free energy and the resulting bound fraction, alongside the longest "
          "complementary run. Per §5.1 this is a ranking input, not a pass/fail gate.")
para(doc, "**The same calculation runs against the two fixed loops.** A loop is a permanent "
          "single-stranded stretch, so unlike the rest of the switch it is always available "
          "to pair with something. If a loop happens to be complementary to one of the "
          "triggers, it becomes an off-pathway sink that no amount of hairpin design will "
          "fix. Scanned against every 30-nt window of our two real transcripts:")
table(doc, ["loop", "worst ΔG against a GFP window", "worst ΔG against an mCherry window",
            "windows below −10 kcal/mol"],
      [["RBS loop, 18 nt", "−11.7", "−13.8", "36 of 730 (GFP), 70 of 2167 (mCherry)"],
       ["Secondary loop, 15 nt", "−9.2", "−7.6", "none"]],
      [3.4, 4.4, 4.8, 4.4], size=8.5)
para(doc, "For scale, a trigger binding its intended footprint runs −30 to −45 kcal/mol, so "
          "neither loop is close to competitive and we do not expect this to decide anything. "
          "It is one extra duplex calculation per design, it is the kind of accident that is "
          "invisible until it bites, and the RBS loop is not entirely clean — so it is "
          "reported per design. Whether it should ever become a filter is decision **D12**.")

# ================================================================= 8
h1(doc, "8.  Validation the code must run")
para(doc, "Each item below is an assertion that halts the run, not a line in a report.")
table(doc, ["#", "assertion"],
      [["1", "Every declared helix pairs antiparallel with zero mismatches (G·U wobbles "
              "allowed only where explicitly declared). This is the check that catches "
              "domain-order errors."],
       ["2", "The intended OFF dot-bracket, written from the element spans and independent "
              "of any folding, is diffed against both the MFE and the centroid."],
       ["3", "The cap sits at the 5' end of the finished construct, not stranded internally."],
       ["4", "(len_main_pre + len_LINKER) mod 3 = 0."],
       ["5", "No in-frame stop codon from the AUG through the end of the GFP CDS."],
       ["6", "No AUG between the RBS loop and the intended start codon that could initiate "
              "out of frame."],
       ["7", "The RBS appears exactly once; no second SD-like motif in the added 5' region."],
       ["8", "MainZ binds k1* strictly weaker than trigger A's k1 does, by at least "
              "ddG_pref. Same for SecondaryZ against k2."],
       ["9", "4 ≤ len_x ≤ 14, and the overlap satisfies §5.1.1: no mismatches at "
              "|x| ≤ 5, at most 2 for |x| = 6–10, at most 3 for |x| = 11–14, never at either "
              "end and never more than 2 consecutive."],
       ["9b", "The measured duplex free energy and bound fraction between the two real "
               "trigger sequences are recorded for every design, together with the longest "
               "complementary run. **No design is rejected on any of the three** — they "
               "rank (§7.4)."],
       ["10", "len_x + len_k2 = 18, and len_main_pre + 3 + len_k1 = 18 — an arm span, "
               "with the bulge counted (§4.1)."],
       ["11", "Trigger A's x and trigger B's Secondary_pre are reverse complements over the "
               "matched positions."],
       ["12", "Any design carried between stages is matched on its full identity "
               "(len_x, len_main_pre, len_k1, len_k2, ddG_pref, mismatch pattern), never on "
               "a partial key."],
       ["13", "The RBS is 3'-flush in the loop, and the SD→AUG spacing equals len_k1 and is "
               "reported for every design (§6.6)."],
       ["14", "Any pattern-prevention list passed to a sequence designer is scoped to "
               "designed positions only. Applied to the whole strand it would reject the "
               "ribosome binding site itself (§6.6)."],
       ["15", "Where the MFE and the centroid disagree on whether the start codon is paired, "
               "the design is flagged in the output rather than silently ranked (§7.1)."],
       ["16", "Each fixed loop is scored against both triggers as a possible off-pathway "
               "duplex, and the margin against the trigger's intended footprint is reported "
               "(§7.4)."]],
      [0.9, 16.1], size=8.5)

doc.add_page_break()

# ================================================================= 9
h1(doc, "9.  Open decisions")
para(doc, "These are the decisions we need before implementation. D1, D2, D5, D8, D9 and D11 "
          "carry a recommendation we are ready to act on unless you disagree; **D3, D4, D6, "
          "D7, D10 and D12 are genuine choices** where we do not think the evidence settles "
          "it.")
table(doc,
      ["#", "decision", "our recommendation", "why it matters"],
      [["D1", "What does \"18\" count — base pairs or arm span?",
        "**18 nt of arm span, bulge included**",
        "So the main arm is 9 + 3 + 6 and only 15 bp are paired. Measured from the published "
        "sequences: Green's main arm is 9 + 3 + 6, Kim's is 11 + 1 + 6, and **both span 18 "
        "nt**. An 18-base-pair reading would give a 21-nt arm, deeper than anything either "
        "group or we have built (§4.1)."],
       ["D2", "Split of the main arm", "**len_k1 = 6, len_main_pre = 9**",
        "len_k1 is the Shine–Dalgarno-to-start-codon spacing and every reference design puts "
        "it at 6; D1 then fixes len_main_pre. The split itself is thermodynamically inert "
        "(−9.61 against −9.57 kcal/mol of drive)."],
       ["D3", "Trigger B total length", "**open** — 50 nt as instructed, but 30–35 is the "
                                        "published precedent",
        "50 nt gives a 32-nt toehold where Kim used 13. Longer buys specificity but makes "
        "trigger B bind far harder than trigger A, which may unbalance the gate."],
       ["D4", "GFP's own leading ATG", "**open** — we lean to deleting it",
        "Green and Kim both keep it; our own switch generator deletes it. Measured: no "
        "thermodynamic difference, and no second start site either way (§6.5). A preference "
        "rather than a result, which is why we would rather you called it."],
       ["D5", "Overlap range and mismatch rule",
        "**|x| = 4–14; 0 mismatches to |x| = 5, 2 for 6–10, 3 for 11–14; never at an end, "
        "never more than 2 in a row**",
        "Every |x| in range is designed and scored. The rule exists to widen the search, and "
        "it does so decisively: for our gene pair there is **no** perfectly complementary "
        "site at any |x| ≥ 9, and the rule finds 359 at |x| = 9. It is a search-space rule, "
        "not a sequestration guarantee — that is measured and ranked separately (§5.1.1)."],
       ["D6", "Size and composition of the lab panel", "**open** — 12 constructs proposed "
                                                       "in §5.3",
        "Determines how many levels of |x| and ddG_pref we can afford to test."],
       ["D7", "Inhibitory loop — fixed or designed?",
        "**open** — Kim's fixed 15-nt loop first, de-novo design where it fails the fold check",
        "Kim used a fixed loop for the series we copy and de-novo NUPACK design for a "
        "different one. Loops are demonstrably not portable between sequence contexts "
        "(§6.6), and the secondary hairpin changes with every gene pair."],
       ["D8", "Include an a = 4 benchmark construct", "**yes**",
        "a = 0 is one step beyond Kim's tested range. A single a = 4 construct puts a "
        "known-working point in the panel."],
       ["D9", "Accessibility metric", "**joint (whole-window) unpaired probability**",
        "The averaged per-base metric does not correspond to what a ribosome requires "
        "(§7.2)."],
       ["D10", "MFE or centroid when they disagree?", "**open** — we lean to the centroid "
                                                      "plus the pairing-probability matrix",
        "At full construct length P(MFE) is about 3 × 10⁻¹⁵ and the two structures disagree "
        "about whether the start codon is paired in 2 of our 5 tested switches (§7.1). The "
        "counter-argument is comparability: the published tables and our own candidates "
        "report the MFE."],
       ["D11", "RBS loop sequence", "**keep the 18-nt loop AGACAAGAACAGAGGAGA**",
        "It is the loop our tested switches used, and it holds the intended fold in all five. "
        "Noted for the record: Green's shorter 11-nt loop measures marginally better on RBS "
        "accessibility, 0.77 against 0.68 (§6.6)."],
       ["D12", "Screen the fixed loops as accidental trigger sites?",
        "**open** — we recommend reporting it, not filtering on it",
        "A loop is permanently single-stranded, so it is the one part of the switch always "
        "free to pair with a trigger. Measured: the RBS loop reaches −13.8 kcal/mol against "
        "real 30-nt windows and the secondary loop stays above −9.2, both far from the "
        "−30 to −45 of a productive trigger duplex (§7.4)."]],
      [0.9, 3.6, 4.6, 7.9], size=8.5,
      row_fills=[None, None, "FFF6E5", "FFF6E5", None, "FFF6E5", "FFF6E5", None, None,
                 "FFF6E5", None, "FFF6E5"])
para(doc, "Table 4 — shaded rows are the six genuine choices. The decision on the input gene "
          "pairs has been dropped from this list at your request; the search still needs "
          "them, it is simply not a design decision.", size=8.5, italic=True, color=GREY)

# ================================================================= Appendix
h1(doc, "Appendix A.  Fixed sequences")
table(doc, ["name", "sequence", "notes"],
      [["cap", "GGG", "3 nt; transcription start, 5' end of the construct"],
       ["RBS", "AACAGAGGAGA", "11 nt, fixed. Identical in Green, in Kim and in our own "
                              "generator. Sits flush against the 3' end of the RBS loop."],
       ["RBS loop", "AGACAAGAACAGAGGAGA", "18 nt = designed flank AGACAAG + the fixed RBS. "
                                          "Equals Kim's 15-nt loop with AGA added, and "
                                          "Green's 11-nt loop with AGACAAG added (§6.6). "
                                          "See D12."],
       ["Secondary loop", "CAAGAACUUAGACAA", "15 nt; the inhibitory-hairpin loop of Kim's "
                                             "Sw-G5-G3n* series (their Table S2-1). Verified "
                                             "to contain no SD-like motif. A bare AUG here "
                                             "would be harmless — see §6.7. Screened as a "
                                             "possible trigger site, D12."],
       ["LINKER", "AACCUGGCGGCAGCGCAAAAG", "21 nt = N L A A A Q K; no in-frame stop. "
                                           "Identical in Green and Kim, and contains no "
                                           "SD-like motif, so it creates no second start "
                                           "site (§6.5)."],
       ["GFP CDS", "ATGCGTAAAGGAGAAGAACTT…  (759 nt)",
        "253 codons; divisible by 3; the only in-frame stops are the terminal TAA TAA. "
        "Begins with its own ATG — see D4."],
       ["prevent list", "aaaa, cccc, gggg, uuuu, kkkkkk, mmmmmm, rrrrrr, ssssss, wwwwww, "
                        "yyyyyy",
        "Kim's pattern-prevention list, identical to the one in our own generator. Applies "
        "to designed positions only — the RBS itself violates rrrrrr (§6.6)."]],
      [2.8, 6.4, 7.8], size=8.5, mono_cols=(1,))

h1(doc, "Appendix B.  Sources and methods")
bullet(doc, "**Kim et al., 2019**, Modulating responses of toehold switches by an "
            "inhibitory hairpin — the two-hairpin architecture, the a / * trade-off, and "
            "the finding that small a produces a genuine AND gate.", size=9.5)
bullet(doc, "**Robson & Green, 2026**, Toehold-VISTA (Nucleic Acids Research) — the "
            "Series A / tsgen2 switch scaffold, the 21-nt linker, and the in-frame-stop "
            "filter.", size=9.5)
bullet(doc, "**Geometry in Table 1** was measured, not quoted: each construct was rebuilt "
            "from the supplementary sequences and folded with ViennaRNA 2.7.2, default "
            "model, 37 °C. Kim's stem lengths were confirmed against the NUPACK design "
            "specification in their supporting information.", size=9.5)
bullet(doc, "**Duplex energies and sequestration** (Figure 3, §5.1) use ViennaRNA "
            "duplexfold at 37 °C, averaged over 400–500 random sequences per point, with "
            "both transcripts at 10 nM. Bound fraction is the solution of A + B ⇌ AB at "
            "equal totals.", size=9.5)
bullet(doc, "**Bulge advantage** (§6.4) and **P(MFE) decay** (§7.1) were computed with "
            "ViennaRNA fold and partition-function routines on the same constructs.", size=9.5)
bullet(doc, "**Loop sequences** were read from the primary sources, not inferred: Green's "
            "from the fixed 5' stem / loop / 3' stem string given in the Toehold-VISTA "
            "methods; Kim's from the switch sequences in their Supplementary Table S2, and "
            "their de-novo design specification from the DU+ script reproduced in their "
            "supporting information. Our own loop and prevent list were read from the lab's "
            "switch generator.", size=9.5)
bullet(doc, "**Loop substitution, RBS and start-codon accessibility, the GFP "
            "internal-initiation scan, in-frame-stop rates, the mismatch placement study and "
            "the MFE-versus-centroid comparison** were computed with ViennaRNA 2.7.2 at "
            "37 °C, default model. Joint unpaired probabilities use a hard-constrained "
            "partition function; bound fractions solve A + B ⇌ AB with both strands at "
            "10 nM.", size=9.5)
bullet(doc, "This document is generated by spec/build_spec.py. Edit that script and "
            "re-run rather than editing the .docx.", size=9.5)

print("wrote", save(doc, OUT))
