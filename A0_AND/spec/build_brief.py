# -*- coding: utf-8 -*-
"""Build A0_AND_Brief.docx - the short version.

Four topics, about a page each: the architecture and its parts, the design
rules, how designs are measured, and the decisions we need.  No background:
where a published design is the reason for a number, it is cited in one line
and nothing more.

    python spec/build_brief.py

Styling is shared with build_spec.py through docx_style.py.
"""
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from docx.shared import Pt
from docx_style import (new_document, save, table, h1, h2, para, bullet, mono,
                        callout, figure, GREY, NAVY, ROOT)

OUT = os.path.join(ROOT, "A0_AND_Brief.docx")

doc = new_document(margin_cm=1.8)

# ------------------------------------------------------------------ title
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("A0 AND Gate")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = NAVY
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(8)
r = p.add_run("Architecture brief — parts, rules, measurements, decisions")
r.font.size = Pt(12)
r.font.color.rgb = GREY

para(doc, "One mRNA that produces GFP only when two RNA inputs are present together. "
          "A main hairpin sequesters the ribosome binding site and the start codon; an "
          "inhibitory hairpin immediately upstream of it sequesters the site where the main "
          "trigger would land. **Trigger B** opens the inhibitory hairpin and exposes that "
          "site; **trigger A** then opens the main hairpin. The variant is named A0 because "
          "the spacer between the two hairpins, a, is **0** — so trigger A has no exposed "
          "nucleotide anywhere on the switch until trigger B has acted.", size=9.5)
para(doc, "This brief states what the construct is. The companion specification carries the "
          "measurements behind each number.", size=9, italic=True, color=GREY, space_after=8)

# ================================================================== 1
h1(doc, "1.  The construct")
figure(doc, "fig1_architecture.png",
       "**Figure 1.** The switch in the OFF state, and the two triggers. Domains that pair "
       "share a colour. **|x| = 9 is drawn only as an example** — len_x is the one free "
       "design parameter, and len_k2 = 18 − len_x follows from it, so both change together "
       "(R10).", width_cm=17.0)

h2(doc, "1.1  Every element, 5' → 3'", new_page=True)
para(doc, "“Pairs with” describes the OFF state. Only **len_x** and **len_r2** are free; "
          "everything else is fixed or derived.", size=9)
table(doc,
      ["#", "element", "pairs with", "length (nt)", "sequence determined by"],
      [["1", "cap", "— free", "3", "fixed: GGG"],
       ["2", "r2*", "— free (trigger-B toehold)", "len_r2 = 32", "revcomp(r2); gene B"],
       ["3", "Secondary_pre*", "x*  (#9)", "len_x", "= x; gene A"],
       ["4", "k2*", "SecondaryZ  (#8)", "18 − len_x", "revcomp(k2); gene B"],
       ["5", "SUSA", "SUSD  (#7)", "0", "placeholder"],
       ["6", "Secondary loop", "— free", "15", "fixed; must carry no RBS-like motif"],
       ["7", "SUSD", "SUSA  (#5)", "0", "placeholder"],
       ["8", "SecondaryZ", "k2*  (#4)", "18 − len_x", "≈ k2, weakened by ddG_pref (R6)"],
       ["9", "x*", "Secondary_pre*  (#3)", "len_x", "revcomp(x); gene A"],
       ["10", "a", "— free", "**0**", "the defining parameter of A0"],
       ["11", "Main_pre*", "Main_pre  (#19)", "**9**", "revcomp(Main_pre); gene A"],
       ["12", "opposing_bulge", "— free in OFF;\npaired by trigger A in ON", "3",
        "revcomp of trigger A's 3 nt (R7)"],
       ["13", "k1*", "MainZ  (#17)", "**6**", "revcomp(k1); gene A"],
       ["14", "MUSA", "MUSD  (#16)", "0", "placeholder"],
       ["15", "RBS loop", "— free", "18", "fixed; RBS flush at its 3' end"],
       ["16", "MUSD", "MUSA  (#14)", "0", "placeholder"],
       ["17", "MainZ", "k1*  (#13)", "**6**", "≈ k1, weakened by ddG_pref (R6)"],
       ["18", "AUG", "— free", "3", "fixed: start codon"],
       ["19", "Main_pre", "Main_pre*  (#11)", "**9**", "gene A"],
       ["20", "LINKER", "— free", "21", "fixed"],
       ["21", "GFP CDS", "— free", "759", "fixed"]],
      [0.7, 3.1, 4.3, 2.5, 6.4], size=7.6)
para(doc, "**len_x is the only free parameter**; len_k2 = 18 − len_x follows from it. "
          "Elements 5, 7, 10, 14 and 16 are zero-length **placeholders**, kept in the layout "
          "so a future version can add an upper stem to either hairpin, or re-open the "
          "inter-hairpin spacer, without renumbering anything.", size=8, italic=True,
     color=GREY)

h2(doc, "1.2  The two triggers")
mono(doc, "Trigger A  (main, gene A)       5'-  k1 - opposing_bulge - Main_pre - a - x  -3'",
     size=8.5, space_after=2)
mono(doc, "Trigger B  (secondary, gene B)  5'-  k2 - Secondary_pre - r2               -3'",
     size=8.5)
para(doc, "**Reading rule.** A trigger binds antiparallel, so its 5' end pairs with the "
          "**3'-most** element of its footprint on the switch. Listing a trigger's domains "
          "in switch order gives the reverse of the correct sequence.", size=9)
para(doc, "**The shared domain.** Secondary_pre = x*, so trigger A carries x and trigger B "
          "carries x*. That overlap is what couples the two halves of the gate — and it is "
          "also what lets the two triggers bind each other instead of the switch (R10).",
     size=9)

h2(doc, "1.3  The four states")
table(doc, ["state", "inputs", "what happens", "output"],
      [["00", "neither", "both hairpins closed; x* is paired", "OFF"],
       ["10", "trigger A only", "x* still paired — A has nothing to nucleate on", "OFF"],
       ["01", "trigger B only", "secondary hairpin opens and frees x*; the AUG stays locked",
        "OFF"],
       ["11", "A + B", "A nucleates on the freed x*, then branch-migrates through Main_pre* "
                       "and k1*, displacing MainZ and Main_pre and freeing the start codon",
        "**ON**"]],
      [1.4, 2.6, 11.0, 2.0], size=8.5,
      row_fills=[None, None, None, "EAF3EC"])


# ================================================================== 2
h1(doc, "2.  Design rules", new_page=True)
table(doc, ["", "rule", "why"],
      [["R1", "**Both hairpin arms span 18 nt**, counting the bulge. Main hairpin: "
              "len_main_pre (9) + bulge (3) + len_k1 (6) = 18, so **15 bp are paired**. "
              "Secondary hairpin: len_x + len_k2 = 18, no bulge, so 18 bp are paired.",
        "Measured from the published sequences: Green's main arm is 9 + 3 + 6 and Kim's is "
        "11 + 1 + 6 — **both span exactly 18 nt**. Counting 18 base pairs instead would give "
        "a 21-nt arm, deeper than anything either group or we have built."],
       ["R2", "**len_k1 = 6**, which with R1 gives **len_main_pre = 9**.",
        "The RBS sits flush against the 3' end of its loop, so len_k1 **is** the "
        "Shine–Dalgarno-to-start-codon spacing. That spacing is 6 nt in Green, in Kim, and "
        "in all five of our own tested switches. 9 + 3 + 6 then reproduces Green's Series A "
        "main hairpin exactly, and equals our own deepest tested candidate."],
       ["R3", "**a = 0.** The two hairpins are directly adjacent.",
        "The defining parameter. Kim's series shows the gate becomes more AND-like as a "
        "shrinks; A0 takes it to the limit, so the dependency is structural rather than a "
        "matter of degree."],
       ["R4", "**Everything between the start codon and the GFP CDS must sum to a multiple "
              "of 3.** Today that is Main_pre (9) + LINKER (21).",
        "Reading frame. Stated in the general form so it still holds if an element is "
        "inserted there later."],
       ["R5", "**No in-frame stop codon in Main_pre.** A hard filter inside the search, not "
              "a check afterwards.",
        "Main_pre is a literal stretch of gene A translated as an N-terminal extension of "
        "GFP. At len_main_pre = 9 this costs 7–10 % of candidate windows on our real "
        "transcripts."],
       ["R6", "**MainZ and SecondaryZ never bind their partner more tightly than the real "
              "trigger does.** A margin ddG_pref weakens them further; it is swept, and "
              "**ddG_pref = 0 (isoenergetic) is one of the levels**, not excluded.",
        "Each Z is a copy of the trigger's own domain. A perfect Watson–Crick copy would "
        "bind harder than the real trigger, which carries wobbles — the trigger would be "
        "climbing uphill to displace the clamp meant to release it."],
       ["R7", "**A 3-nt opposing_bulge faces the AUG across the main stem**, forming a "
              "3 × 3 internal loop.",
        "In ON, trigger A pairs straight through it: 18 contiguous base pairs where the "
        "hairpin managed 15 plus an internal loop. Measured advantage to trigger A: "
        "9.6 kcal/mol, against 1.9 with no bulge and 6.5 with Kim's 1-nt bulge. It also "
        "breaks the stem into two shorter helices, reducing RNase III exposure."],
       ["R8", "**The RBS loop is 18 nt: a designed 7-nt flank followed by the fixed 11-nt "
              "RBS**, which stays flush at the loop's 3' end.",
        "Both reference designs are built the same way and differ only in flank length "
        "(Green 0 nt, Kim 4 nt). The 3'-flush RBS is what makes R2 true."],
       ["R9", "**The secondary loop carries no SD-like motif.**",
        "It must not create a second ribosome entry point. A bare AUG there is harmless — "
        "it sits in the 5'UTR and bacterial initiation needs a Shine–Dalgarno sequence — so "
        "no constraint is placed on start codons in this loop. Kim's 15-nt inhibitory loop "
        "satisfies the SD condition and is our starting sequence."],
       ["R10", "**Overlap: |x| = 4 to 14, every value designed and scored** — not only "
               "the longest available. **Mismatches are allowed at every length, in any "
               "number and any placement.** No design is rejected on a mismatch ratio or a "
               "run-length cap; instead the trigger-trigger duplex energy is computed from "
               "the real sequences for every design and used to rank (§3.4).",
        "|x| ≥ 4 because the overlap is the mechanism; |x| ≤ 14 so that k2 = 18 − |x| "
        "stays at least 4 bp. We deliberately impose nothing else here: neither a 1:4 "
        "mismatch ratio nor a run cap actually bounds sequestration (a 1:4 ratio with "
        "uncontrolled placement is already 74 % sequestered at |x| = 14), and a hard filter "
        "would throw away workable designs — Kim's own working trigger pair sits at 73 %."],
       ["R11", "**Trigger lengths follow: A = 18 + len_x, B = 18 + len_r2.**",
        "A is 30 nt at |x| = 12, the length both reference designs used. B is 50 nt at "
        "len_r2 = 32."],
       ["R12", "**Any pattern-prevention list given to a sequence designer is scoped to "
               "designed positions only.**",
        "The standard list forbids rrrrrr — six purines — and the RBS ends in AGGAGA. "
        "Applied to the whole strand it would reject the ribosome binding site itself."]],
      [1.1, 6.5, 9.4], size=8)

# ================================================================== 3
h1(doc, "3.  How designs are measured", new_page=True)

h2(doc, "3.1  The four states")
para(doc, "Each state is evaluated as a multi-strand equilibrium. Per state we record the "
          "ensemble free energy, the MFE structure **and its probability**, the centroid "
          "structure with its distance from the MFE, the base-pairing probability matrix, "
          "and the accessibility series of §3.2.", size=9.5)

h2(doc, "3.2  Accessibility around the start codon")
para(doc, "The reported quantity is the probability that a defined window is **entirely "
          "unpaired at once**, not the average over its bases:", size=9.5)
mono(doc, "P_open(W)   =  Z( all of W unpaired ) / Z          "
          "dG_open(W)  =  - RT · ln P_open(W)", size=8.5)
para(doc, "A ribosome does not sample bases independently — it needs its whole footprint "
          "free at the same moment. For a window whose bases are each 90 % unpaired, the "
          "averaged metric always reads 0.90 while the joint probability lies anywhere "
          "between 0.90 and far less. ViennaRNA computes the joint form directly: one "
          "hard-constrained partition function per window.", size=9.5)
para(doc, "Windows are nested and all anchored on the start codon: the AUG alone; the RBS "
          "through the AUG; AUG through AUG + 3n for n = 1…10 codons; and the 30-nt ribosome "
          "footprint. The series shows how far the opening extends — which separates a "
          "switch that merely exposes the AUG from one that clears the ribosome's path.",
     size=9.5)

h2(doc, "3.3  Discrimination")
mono(doc, "ddG_AND  =  [ dG_open(11) - dG_open(01) ]  -  [ dG_open(10) - dG_open(00) ]",
     size=8.5)
para(doc, "How much more trigger A opens the switch when trigger B is present than when it "
          "is absent. Length cancels in the difference, so this needs no normalisation. "
          "Separation is reported threshold-free: state 11 against the worst of 00, 01, 10.",
     size=9.5)

h2(doc, "3.4  Trigger–trigger sequestration")
para(doc, "Reported at a stated concentration (10 nM per transcript) as the duplex free "
          "energy, the resulting bound fraction, and the longest complementary run. This is "
          "a **ranking input, not a veto**: Kim's own published trigger pair, in the "
          "construct they describe as a working two-input AND gate, sits at 73 % "
          "sequestration by this calculation.", size=9.5)
para(doc, "**The same check runs against the two fixed loops**, which are permanently "
          "single-stranded and so are the one part of the switch always free to pair with a "
          "trigger. Against every 30-nt window of our real transcripts the secondary loop "
          "stays above −9.2 kcal/mol; the RBS loop reaches −13.8, with 36 of 730 GFP "
          "windows below −10. Both are far from the −30 to −45 of a productive trigger "
          "duplex, so this is reported per design, not filtered on (D12).", size=9.5)

h2(doc, "3.5  Which structure the ranking reads")
para(doc, "P(MFE) collapses with length. On our own switches it is 0.02–0.08 for the switch "
          "alone and about 3 × 10⁻¹⁵ for the full construct with GFP — and at that length the "
          "MFE and the centroid **disagree about whether the start codon is paired in 2 of "
          "5 constructs**. The ranking therefore reads the centroid and the pairing "
          "probability matrix; the MFE is reported but trusted only over the isolated "
          "hairpin region. Disagreements are flagged, not silently ranked. (Decision D11.)",
     size=9.5)

h2(doc, "3.6  Assertions that halt the run")
para(doc, "Each of these stops the run rather than appearing as a line in a report. "
          "**(1)** every declared helix pairs antiparallel with zero mismatches, wobbles only "
          "where declared — this is what catches domain-order errors; "
          "**(2)** the intended OFF dot-bracket, written from the element spans of §1.1 and "
          "independent of any folding, is diffed against both the MFE and the centroid; "
          "**(3)** (len_main_pre + len_LINKER) mod 3 = 0, and no in-frame stop from the AUG "
          "through the end of the GFP CDS; "
          "**(4)** the RBS appears exactly once, flush at the loop's 3' end, and the "
          "SD→AUG spacing equals len_k1 and is reported; "
          "**(5)** no AUG between the RBS loop and the intended start codon that could "
          "initiate out of frame; "
          "**(6)** MainZ never binds k1* more tightly than trigger A's k1 does, and "
          "SecondaryZ likewise against k2; "
          "**(7)** 4 ≤ len_x ≤ 14; len_x + len_k2 = 18; len_main_pre + 3 + len_k1 = 18; "
          "**(8)** trigger A's x and trigger B's Secondary_pre are reverse complements over "
          "the matched positions; "
          "**(9)** the cap sits at the 5' end of the finished construct, not stranded "
          "internally; "
          "**(10)** any design carried between stages is matched on its full identity — "
          "len_x, len_main_pre, len_k1, len_k2, ddG_pref, mismatch pattern — never a "
          "partial key.", size=9)

# ================================================================== 4
h1(doc, "4.  Decisions we need", new_page=True)
para(doc, "Six carry a recommendation we are ready to act on unless you disagree. The six "
          "shaded rows are genuine choices where we do not think the evidence settles it. "
          "**D1 is the one that changed most recently** — it was previously written as 18 "
          "base pairs, which the published sequences do not support.", size=9.5)
table(doc, ["#", "decision", "our position"],
      [["D1", "What does \"18\" count?", "**18 nt of arm span, bulge included** — so the main "
                                       "hairpin is 9 + 3 + 6 and only **15 bp are paired**. "
                                       "Measured: Green's main arm is 9 + 3 + 6 and Kim's is "
                                       "11 + 1 + 6, both spanning 18 nt. Reading it as 18 "
                                       "base pairs would give a 21-nt arm, deeper than "
                                       "anything published or tested."],
       ["D2", "Split of the main arm", "**len_k1 = 6, len_main_pre = 9.** len_k1 is the "
                                       "SD→AUG spacing and every reference design puts it at "
                                       "6; R1 then fixes len_main_pre. The split itself is "
                                       "thermodynamically inert: −9.61 kcal/mol of drive at "
                                       "9 + 3 + 6 against −9.57 at 12 + 3 + 6."],
       ["D3", "Trigger B total length", "**Open.** 50 nt as instructed gives a 32-nt "
                                        "toehold where Kim used 13. Longer buys specificity "
                                        "but makes B bind far harder than A, which may "
                                        "unbalance the gate. 30–35 nt is the precedent."],
       ["D4", "GFP's own leading ATG", "**Open, we lean to deleting it.** Green and Kim keep "
                                       "it; our own switch generator deletes it. Measured: "
                                       "no thermodynamic difference and no second start site "
                                       "either way. A preference, not a result."],
       ["D5", "Overlap range and mismatch rule", "**|x| = 4–14; mismatches unrestricted; "
                                                 "rank on the measured trigger-trigger "
                                                 "duplex, reject nothing.** Neither a 1:4 "
                                                 "ratio nor a run cap actually bounds "
                                                 "sequestration, and a hard filter would "
                                                 "discard workable designs."],
       ["D6", "Size and composition of the lab panel", "**Open.** 12 constructs proposed: "
                                                       "8 gate designs (4 levels of |x| × 2 "
                                                       "of ddG_pref), 1 benchmark at a = 4, "
                                                       "3 controls."],
       ["D7", "Inhibitory loop — fixed or designed?", "**Open.** Start from Kim's fixed "
                                                     "15-nt loop and design de novo where it "
                                                     "fails the fold check. Loops are not "
                                                     "portable: substituting one into our "
                                                     "five tested switches refolded all five."],
       ["D8", "Include an a = 4 benchmark construct", "**Yes.** a = 0 is one step beyond "
                                                      "Kim's shortest tested spacing; one "
                                                      "a = 4 construct puts a known-working "
                                                      "point in the panel."],
       ["D9", "Accessibility metric", "**Joint whole-window unpaired probability** (§3.2). "
                                      "The averaged per-base metric does not correspond to "
                                      "what a ribosome requires."],
       ["D10", "MFE or centroid when they disagree?", "**Open, we lean to the centroid** plus "
                                                      "the pairing-probability matrix (§3.5). "
                                                      "The counter-argument is comparability: "
                                                      "published tables report the MFE."],
       ["D11", "RBS loop sequence", "**Keep the 18-nt loop AGACAAGAACAGAGGAGA.** It holds "
                                    "the intended fold in all five tested switches. Noted "
                                    "for the record: Green's shorter 11-nt loop measures "
                                    "marginally better on RBS accessibility, 0.77 to 0.68."],
       ["D12", "Screen the loops as accidental trigger sites?",
        "**Open, we recommend reporting it and not filtering on it.** The RBS loop reaches "
        "−13.8 kcal/mol against real 30-nt windows (36 of 730 GFP windows below −10); the "
        "secondary loop stays above −9.2. Both are far from the −30 to −45 a trigger gets "
        "on its intended footprint, so this is a diagnostic, not an expected failure."]],
      [0.9, 4.0, 12.1], size=8,
      row_fills=[None, None, "FFF6E5", "FFF6E5", None, "FFF6E5", "FFF6E5", None, None,
                 "FFF6E5", None, "FFF6E5"])

# ================================================================== Appendix
h1(doc, "Appendix.  Fixed sequences")
table(doc, ["name", "sequence", "notes"],
      [["cap", "GGG", "3 nt; transcription start, 5' end of the construct"],
       ["RBS", "AACAGAGGAGA", "11 nt, fixed. Identical in Green, in Kim and in our own "
                              "generator."],
       ["RBS loop", "AGACAAGAACAGAGGAGA", "18 nt = designed flank AGACAAG + the fixed RBS, "
                                          "which stays flush at the 3' end. See D12."],
       ["Secondary loop", "CAAGAACUUAGACAA", "15 nt; the inhibitory-hairpin loop of Kim's "
                                             "Sw-G5-G3n* series. Verified to carry no "
                                             "SD-like motif. See D7."],
       ["LINKER", "AACCUGGCGGCAGCGCAAAAG", "21 nt = N L A A A Q K. No in-frame stop, and no "
                                           "SD-like motif, so it creates no second start "
                                           "site."],
       ["GFP CDS", "ATGCGTAAAGGAGAAGAACTT… (759 nt)",
        "253 codons, divisible by 3; the only in-frame stops are the terminal TAA TAA. "
        "Begins with its own ATG — see D4."],
       ["prevent list", "aaaa, cccc, gggg, uuuu, kkkkkk, mmmmmm, rrrrrr, ssssss, wwwwww, "
                        "yyyyyy",
        "Kim's pattern-prevention list, identical to the one in our own generator. Applies "
        "to designed positions only (R12)."]],
      [2.6, 6.2, 8.2], size=8, mono_cols=(1,))

para(doc, "Sources: Kim et al. 2019, Modulating responses of toehold switches by an "
          "inhibitory hairpin — the two-hairpin architecture and the a trade-off. "
          "Robson & Green 2026, Toehold-VISTA (NAR) — the switch scaffold, the 21-nt linker "
          "and the in-frame-stop filter. All geometry and every energy quoted here was "
          "measured with ViennaRNA 2.7.2 at 37 °C, default model; see the full "
          "specification for method detail.", size=8.5, italic=True, color=GREY)

print("wrote", save(doc, OUT))
